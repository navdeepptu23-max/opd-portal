import os
import csv
import zipfile
from io import BytesIO, StringIO
from textwrap import wrap
from datetime import datetime, timedelta
from flask import Flask, render_template, redirect, url_for, flash, request, abort, session, Response, jsonify
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from flask_wtf import FlaskForm
from flask_wtf.csrf import CSRFProtect
from wtforms import StringField, PasswordField, SelectField, SubmitField
from wtforms.validators import DataRequired, Length, EqualTo, ValidationError
from werkzeug.security import generate_password_hash, check_password_hash
from reportlab.lib.pagesizes import A4, landscape
from reportlab.pdfgen import canvas
from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from modules.morbidity import morbidity_bp

app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'change-this-in-production-use-env-var')

database_url = os.environ.get('DATABASE_URL', '').strip()
if database_url.startswith('postgres://'):
    # Render/Heroku may provide postgres://, but SQLAlchemy expects postgresql://
    database_url = database_url.replace('postgres://', 'postgresql://', 1)
if not database_url:
    # On Render (RENDER env is set) with no DATABASE_URL → the free DB likely
    # expired. Use SQLite as emergency read-only fallback so the app boots and
    # shows an admin warning rather than crashing entirely.
    # Locally (no RENDER env) → SQLite is fine for development.
    database_url = 'sqlite:///portal.db'
    print('[STARTUP] WARNING: DATABASE_URL is not set – using SQLite fallback.')
    print('[STARTUP] If on Render, your free PostgreSQL likely expired.')
    print('[STARTUP] Create a new database and update DATABASE_URL in Environment settings.')

app.config['SQLALCHEMY_DATABASE_URI'] = database_url
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['WTF_CSRF_TIME_LIMIT'] = 3600
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(hours=12)
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
app.config['SESSION_DEBUG'] = os.environ.get('SESSION_DEBUG', '0') == '1'

db = SQLAlchemy(app)
csrf = CSRFProtect(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'
login_manager.login_message = 'Please log in to access this page.'
login_manager.login_message_category = 'warning'
login_manager.session_protection = None
app.register_blueprint(morbidity_bp)


@app.before_request
def _session_debug_before_request():
    if not app.config.get('SESSION_DEBUG'):
        return
    sid = request.cookies.get(app.config.get('SESSION_COOKIE_NAME', 'session'))
    sid_preview = sid[:16] + '...' if sid else None
    app.logger.info(
        'SESSION_DEBUG before path=%s method=%s auth=%s user_id=%s remember=%s sid=%s',
        request.path,
        request.method,
        current_user.is_authenticated,
        session.get('_user_id'),
        session.get('_remember'),
        sid_preview,
    )


_USING_SQLITE = 'sqlite' in database_url.lower()


@app.before_request
def _sqlite_warning():
    """Flash a prominent warning when running on SQLite (database likely expired)."""
    if not _USING_SQLITE:
        return
    if not current_user.is_authenticated:
        return
    if request.endpoint in ('static', 'login', 'logout'):
        return
    if not session.get('_sqlite_warned'):
        flash(
            'WARNING: The PostgreSQL database is not connected. '
            'Data entered now will be LOST on next restart. '
            'Admin must create a new database on Render and update DATABASE_URL.',
            'danger',
        )
        session['_sqlite_warned'] = True


@app.before_request
def _enforce_active_account_session():
    # Prevent stale remember-cookie sessions for deactivated accounts.
    # Restrict this to navigational requests so form POST submissions are not
    # interrupted by session churn in edge proxy/browser cases.
    if request.method != 'GET':
        return
    if not current_user.is_authenticated:
        return
    if current_user.is_active:
        return
    if request.endpoint in ('logout', 'login', 'static'):
        return
    logout_user()
    session.clear()
    flash('Your account is inactive. Contact admin.', 'danger')
    return redirect(url_for('login'))


@app.after_request
def _session_debug_after_request(response):
    if app.config.get('SESSION_DEBUG'):
        app.logger.info(
            'SESSION_DEBUG after path=%s status=%s set_cookie=%s',
            request.path,
            response.status_code,
            'Set-Cookie' in response.headers,
        )
    return response


HOSPITAL_INDICATOR_DEFAULTS = [
    (1, 'Total OPD Attendance'),
    (2, 'Total IPD Admissions'),
    (3, 'Emergency Cases Handled'),
    (4, 'Maternal Cases Registered'),
    (5, 'Institutional Deliveries'),
    (6, 'Immunization Sessions Conducted'),
    (7, 'Lab Tests Performed'),
    (8, 'Referral Cases Sent'),
    (9, 'Referral Cases Received'),
    (10, 'Road Traffic Accident Cases'),
]

PROFORMA_HPI_DEFAULTS = [
    ('1A', 'NO. OF OUTPATIENTS: NEW'),
    ('1B', 'OLD'),
    ('1C', 'EMERGENCY'),
    ('1D', 'TOTAL'),
    ('2', 'NO OF ADMISSIONS DURING THE MONTH'),
    ('3', 'NO. OF ADMISSIONS THROUGH EMERGENCY (OUT OF 2)'),
    ('4', 'NO. OF MEDICAL LEGAL CASES ADMITTED (OUT OF 2)'),
    ('5', 'NO. OF PATIENTS ADMITTED AND DISCHARGED ON THE SAME DAY (OUT OF 2)'),
    ('6', 'NO. OF TUBECTOMIES INCLUDING LAPAROSCOPIC'),
    ('7', 'NO OF VASECTOMIES'),
    ('8', 'NO. OF MINOR SURGERIES (EXCLUDING VASECTOMIES)'),
    ('9', 'NO. OF MAJOR SURGERIES (EXCLUDING TUBECTOMIES)'),
    ('10', 'TOTAL NO. OF SURGERIES (5+6+7+8)'),
    ('11', 'NO. OF DEATHS (Please mention no of maternal and infant hospital deaths in remarks column separately)'),
    ('13', 'NO. OF NORMAL DELIVERIES'),
    ('14', 'NO. OF CAESAREAN DELIVERIES'),
    ('15', 'TOTAL NO. OF DELIVERIES (13+14)'),
    ('16', 'MALE CH (EXCLUDING STILL BIRTH OUT OF 15)'),
    ('17', 'FEMALE CH (EXCLUDING STILL BIRTH OUT OF 15)'),
    ('21', 'NO. OF LAB-TESTS'),
    ('22', 'TOTAL NO. OF CUMULATIVE INPATIENTS DAYS'),
    ('23', 'User Charges Collection during the month (in Rs.)'),
    ('24', 'Number of RSBY Cases during the month'),
]

PROFORMA_II_DEFAULTS = [
    (2, 'Typhoid Fever and Paratyphoid Fever'),
    (5, 'Amoebiasis'),
    (6, 'Diarrhoea'),
    (8, 'Respiratory TB'),
    (10, 'T.B of other organ'),
    (29, 'Measles'),
    (31, 'Other Viral Hepatitis'),
    (32, 'HIV'),
    (37, 'Helminthiasis'),
    (79, 'Other Anaemia'),
    (85, 'Disorders of thyroid glands'),
    (86, 'Diabetes Mellitus'),
    (89, 'Mental and behavioural Disorder'),
    (98, 'Diseases of eye'),
    (99, 'Diseases of the ear'),
    (102, 'Hypertensive Heart Disease'),
    (103, 'All other hypertensive diseases'),
    (114, 'Pharyngitis & Tonsillitis'),
    (116, 'Other Acute Upper Respiratory Infections'),
    (118, 'Acute Bronchitis'),
    (119, 'Ch. Bronchitis and unspecified Emphysema'),
    (120, 'Asthma'),
    (121, 'Other lower respiratory disorders'),
    (126, 'Diseases of Oral Cavity'),
    (127, 'Gastric And Duodenal ulcer'),
    (128, 'Gastritis & Duodenitis'),
    (134, 'Cholelithiasis And Cholecystitis'),
    (136, 'Other Diseases other part of Digestive system'),
    (137, 'Infections of Skin'),
    (138, 'All other disease of Skin'),
    (139, 'Rheumatoid Arthritis & other inflammatory Polyarthropathies'),
    (147, 'Other diseases of Urinary Track'),
    (149, 'All other Diseases of male genital organs'),
    (151, 'All other diseases of female genital organs'),
    (152, 'Spontaneous Abortion'),
    (155, 'Oedema/ Proteinuria & Hypertension Disorder in Pregnancy/childbirth & puerperium'),
    (157, 'Obstructed Labour'),
    (158, 'Complication predominantly related to puerperium'),
    (161, 'All other obstetric conditions not elsewhere classified'),
    (172, 'Abdominal and Pelvic pain'),
    (175, 'Fever of Unknown origin (PUO)'),
    (180, 'All other Symptoms, Signs & abnormal clinical lab findings not elsewhere classified'),
    (186, 'Dislocations, sprains & Strains of body regions'),
    (189, 'Other injuries'),
    (191, 'Burns & Corrosions'),
    (192, 'Poisoning by drugs & Biological substances and toxic effect of substances'),
    (193, 'Other specified effects of external causes & certain early complications of trauma'),
    (198, 'Other Road Side Accidents (RSA)'),
    (215, 'Bites of snake & other Venomous animals/DOG BITE'),
]

CBHI_FORM1_DEFAULTS = [
    (1, 'ACUTE DIARRHOEAL DISEASES (INCLUDING GASTRO ENTERITIS ETC) A09'),
    (2, 'ACUTE POLIOMYELITIS A80'),
    (3, 'ACUTE RESPIRATORY INFECTION (INCLUDING INFLUENZA AND EXCLUDING PNEUMONIA) J00-06, J11.1'),
    (4, 'AIDS (AS REPORTED BY NACO)'),
    (5, 'CHICKEN POX B01'),
    (6, 'CHOLERA (LAB CONFIRMED) A00'),
    (7, 'CORONA'),
    (8, 'DIPHTHERIA (LAB CONFIRMED) A36'),
    (9, 'DRY EYE DISEASE (DED)'),
    (10, 'ENCEPHALITIS G04.0'),
    (11, 'ENTERIC FEVER (LAB CONFIRMED) A01'),
    (12, 'GONOCOCCAL INFECTION A54'),
    (13, 'LEPTOSPIROSIS (LAB CONFIRMED)'),
    (14, 'MEASLES (LAB CONFIRMED) B05'),
    (15, 'MENINGITIS (OTHER THAN BACTERIAL)'),
    (16, 'MENINGOCOCCAL MENINGITIS (LAB CONFIRMED)'),
    (17, 'NEONATAL TETANUS (LAB CONFIRMED) A33'),
    (18, 'OTHER STD DISEASES'),
    (19, 'PLAGUE A20'),
    (20, 'PNEUMONIA J12-18'),
    (21, 'PULMONARY TUBERCULOSIS'),
    (22, 'RABIES A82'),
    (23, 'SCRUB TYPHUS (LAB CONFIRMED)'),
    (24, 'SWINE FLU'),
    (25, 'SYPHILIS A50-A53'),
    (26, 'TETANUS OTHER THAN NEONATAL A35'),
    (27, 'VIRAL HEPATITIS - A (LAB CONFIRMED) B15.9'),
    (28, 'VIRAL HEPATITIS - B (LAB CONFIRMED) B16.9'),
    (29, 'VIRAL HEPATITIS - C (LAB CONFIRMED)'),
    (30, 'VIRAL HEPATITIS - D (LAB CONFIRMED)'),
    (31, 'VIRAL HEPATITIS - E (LAB CONFIRMED)'),
    (32, 'VIRAL HEPATITIS ALL'),
    (33, 'VIRAL MENINGITIS - G03.9'),
    (34, 'WHOOPING COUGH (LAB CONFIRMED) A37'),
]

CBHI_FORM2_DEFAULTS = [
    (1, 'ACCIDENTAL INJURIES S00-S99, T00-T14'),
    (2, 'ARSENICOSIS'),
    (3, 'ASTHMA J45'),
    (4, 'AUTOIMMUNE DISEASES'),
    (5, 'BRONCHITIS J40'),
    (6, 'BURNS'),
    (7, 'CANCER [ORAL, LIP, ORAL CAVITY AND PHARYNX] C00-C14, D10'),
    (8, 'CANCER [BREAST] C50, D24'),
    (9, 'CANCER [CERVIX] C53, D26'),
    (10, 'CANCER [LUNG] C34, D14.3'),
    (11, 'CANCER [OTHERS EXCLUDING S. NO. 7 TO 10] C00-D48'),
    (12, 'CEREBROVASCULAR ACCIDENT I60-I69'),
    (13, 'DIABETES MELLITUS [TYPE 1] E10'),
    (14, 'DIABETES MELLITUS [TYPE 2] E11'),
    (15, 'EMPHYSEMAS J43'),
    (16, 'HEART DISEASES [CONGENITAL] Q20-Q28'),
    (17, 'HEART DISEASES [ISCHEMIC] I20-I25'),
    (18, 'HYPERTENSION I10-I15'),
    (19, 'MENTAL DISORDERS F10-F19, F99'),
    (20, 'NEUROLOGICAL DISORDER [CHRONIC] G90-G99'),
    (21, 'NEUROLOGICAL DISORDERS [OTHER EXCLUDING S. NO. 19] F00-F03, G00-G83'),
    (22, 'OBESITY E66.9'),
    (23, 'OTHER CARDIOVASCULAR DISEASES I05-I09, I26-I52, I70-I99'),
    (24, 'OTHERS'),
    (25, 'RENAL FAILURE [ACUTE] N17'),
    (26, 'RENAL FAILURE [CHRONIC] N18'),
    (27, 'RHEUMATIC FEVER I00-I02'),
    (28, 'ROAD TRAFFIC ACCIDENTS V01-V89'),
    (29, 'RUBELLA'),
    (30, 'SEVERE MENTAL DISORDER F99'),
    (31, 'SNAKE BITE T63.0'),
]

PROFORMA_HPI_ORDER = {code: index for index, (code, _) in enumerate(PROFORMA_HPI_DEFAULTS)}


# ── Models ──────────────────────────────────────────────────────────────────

class User(UserMixin, db.Model):
    __tablename__ = 'users'
    id          = db.Column(db.Integer, primary_key=True)
    username    = db.Column(db.String(80), unique=True, nullable=False)
    password_hash = db.Column(db.String(256), nullable=False)
    role        = db.Column(db.String(20), nullable=False, default='sub')
    created_by  = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=True)
    created_at  = db.Column(db.DateTime, default=datetime.utcnow)
    is_active   = db.Column(db.Boolean, default=True)

    creator     = db.relationship('User', remote_side=[id], backref='created_users')

    def set_password(self, password):
        self.password_hash = generate_password_hash(password)

    def check_password(self, password):
        return check_password_hash(self.password_hash, password)

    @property
    def is_super_admin(self):
        return self.role == 'super_admin'

    @property
    def is_admin_or_above(self):
        return self.role in ('super_admin', 'admin')

    def can_manage(self, target_user):
        """Return True if this user has permission to manage target_user."""
        if self.is_super_admin:
            return target_user.id != self.id
        if self.role == 'admin':
            return target_user.role == 'sub'
        return False


class LoginAudit(db.Model):
    __tablename__ = 'login_audits'

    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), nullable=False, index=True)
    user_id = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=True, index=True)
    success = db.Column(db.Boolean, nullable=False, default=False, index=True)
    reason = db.Column(db.String(120), nullable=False, default='')
    ip_address = db.Column(db.String(64), nullable=False, default='')
    user_agent = db.Column(db.String(260), nullable=False, default='')
    created_at = db.Column(db.DateTime, default=datetime.utcnow, index=True)

    user = db.relationship('User', foreign_keys=[user_id])


class HospitalIndicator(db.Model):
    __tablename__ = 'hospital_indicators'

    id = db.Column(db.Integer, primary_key=True)
    month_year = db.Column(db.String(20), nullable=False, index=True)
    indicator_no = db.Column(db.Integer, nullable=False)
    indicator_name = db.Column(db.String(180), nullable=False)
    opd_count = db.Column(db.Integer, nullable=False, default=0)
    ipd_count = db.Column(db.Integer, nullable=False, default=0)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

    __table_args__ = (
        db.UniqueConstraint('month_year', 'indicator_no', name='uq_hospital_indicator_month_no'),
    )


class HospitalIndicatorMeta(db.Model):
    __tablename__ = 'hospital_indicator_meta'

    id = db.Column(db.Integer, primary_key=True)
    month_year = db.Column(db.String(20), nullable=False, unique=True, index=True)
    institution_name = db.Column(db.String(150), nullable=False, default='PHC POSSI')


class ProformaHPIRow(db.Model):
    __tablename__ = 'proforma_hpi_rows'

    id = db.Column(db.Integer, primary_key=True)
    month_year = db.Column(db.String(20), nullable=False, index=True)
    indicator_code = db.Column(db.String(8), nullable=False)
    indicator_label = db.Column(db.String(260), nullable=False)
    male = db.Column(db.Integer, nullable=False, default=0)
    female = db.Column(db.Integer, nullable=False, default=0)
    male_child_u14 = db.Column(db.Integer, nullable=False, default=0)
    female_child_u14 = db.Column(db.Integer, nullable=False, default=0)
    total = db.Column(db.Integer, nullable=False, default=0)
    remarks = db.Column(db.String(260), nullable=False, default='')
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

    __table_args__ = (
        db.UniqueConstraint('month_year', 'indicator_code', name='uq_proforma_hpi_month_code'),
    )


class ProformaHPIMeta(db.Model):
    __tablename__ = 'proforma_hpi_meta'

    id = db.Column(db.Integer, primary_key=True)
    month_year = db.Column(db.String(20), nullable=False, unique=True, index=True)
    hospital_name = db.Column(db.String(200), nullable=False, default='COMPILED REPORT OF BLOCK- POSSI')
    district = db.Column(db.String(120), nullable=False, default='HOSHIARPUR')
    sanctioned_beds = db.Column(db.String(60), nullable=False, default='')
    functional_beds = db.Column(db.String(60), nullable=False, default='')
    doctor_incharge = db.Column(db.String(120), nullable=False, default='SMO POSSI')
    note_text = db.Column(
        db.String(400),
        nullable=False,
        default='Comments of SMO incharge regarding change of functional beds if any, nil reports and unusually large or small comparative figures regarding performance of any department of the hospital should be mentioned in the remarks column.'
    )


class ProformaIIRow(db.Model):
    __tablename__ = 'proforma_ii_rows'

    id = db.Column(db.Integer, primary_key=True)
    month_year = db.Column(db.String(20), nullable=False, index=True)
    sr_no = db.Column(db.Integer, nullable=False)
    disease_name = db.Column(db.String(260), nullable=False)
    opd_count = db.Column(db.Integer, nullable=False, default=0)
    ipd_count = db.Column(db.Integer, nullable=False, default=0)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

    __table_args__ = (
        db.UniqueConstraint('month_year', 'sr_no', name='uq_proforma_ii_month_sr'),
    )


class ProformaIIMeta(db.Model):
    __tablename__ = 'proforma_ii_meta'

    id = db.Column(db.Integer, primary_key=True)
    month_year = db.Column(db.String(20), nullable=False, unique=True, index=True)
    institution_name = db.Column(db.String(200), nullable=False, default='PHC POSSI')


class CbhiForm1Row(db.Model):
    __tablename__ = 'cbhi_form1_rows'

    id = db.Column(db.Integer, primary_key=True)
    month_year = db.Column(db.String(20), nullable=False, index=True)
    sr_no = db.Column(db.Integer, nullable=False)
    disease_name = db.Column(db.String(320), nullable=False)
    code = db.Column(db.String(40), nullable=False, default='')

    general_m = db.Column(db.Integer, nullable=False, default=0)
    general_f = db.Column(db.Integer, nullable=False, default=0)
    general_tr = db.Column(db.Integer, nullable=False, default=0)
    general_total = db.Column(db.Integer, nullable=False, default=0)

    emergency_m = db.Column(db.Integer, nullable=False, default=0)
    emergency_f = db.Column(db.Integer, nullable=False, default=0)
    emergency_tr = db.Column(db.Integer, nullable=False, default=0)
    emergency_total = db.Column(db.Integer, nullable=False, default=0)

    ipd_general_m = db.Column(db.Integer, nullable=False, default=0)
    ipd_general_f = db.Column(db.Integer, nullable=False, default=0)
    ipd_general_tr = db.Column(db.Integer, nullable=False, default=0)
    ipd_general_total = db.Column(db.Integer, nullable=False, default=0)

    ipd_emergency_m = db.Column(db.Integer, nullable=False, default=0)
    ipd_emergency_f = db.Column(db.Integer, nullable=False, default=0)
    ipd_emergency_tr = db.Column(db.Integer, nullable=False, default=0)
    ipd_emergency_total = db.Column(db.Integer, nullable=False, default=0)

    overall_m = db.Column(db.Integer, nullable=False, default=0)
    overall_f = db.Column(db.Integer, nullable=False, default=0)
    overall_tr = db.Column(db.Integer, nullable=False, default=0)
    overall_total = db.Column(db.Integer, nullable=False, default=0)

    deaths_m = db.Column(db.Integer, nullable=False, default=0)
    deaths_f = db.Column(db.Integer, nullable=False, default=0)
    deaths_tr = db.Column(db.Integer, nullable=False, default=0)
    deaths_total = db.Column(db.Integer, nullable=False, default=0)

    remarks = db.Column(db.String(220), nullable=False, default='')
    is_custom = db.Column(db.Boolean, nullable=False, default=False)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

    __table_args__ = (
        db.UniqueConstraint('month_year', 'sr_no', name='uq_cbhi_form1_month_sr'),
    )


class CbhiForm1Meta(db.Model):
    __tablename__ = 'cbhi_form1_meta'

    id = db.Column(db.Integer, primary_key=True)
    month_year = db.Column(db.String(20), nullable=False, unique=True, index=True)
    health_establishment = db.Column(db.String(220), nullable=False, default='PRIMARY HEALTH CENTRE, POSSI')
    complete_address = db.Column(db.String(320), nullable=False, default='VPO: POSSI, TEHSIL-GARHSHANKAR, DISTRICT-HOSHIARPUR (PUNJAB)')
    approving_authority = db.Column(db.String(160), nullable=False, default='')
    authority_designation = db.Column(db.String(160), nullable=False, default='')
    official_email = db.Column(db.String(160), nullable=False, default='')
    official_phone = db.Column(db.String(40), nullable=False, default='')


class CbhiForm2Row(db.Model):
    __tablename__ = 'cbhi_form2_rows'

    id = db.Column(db.Integer, primary_key=True)
    month_year = db.Column(db.String(20), nullable=False, index=True)
    sr_no = db.Column(db.Integer, nullable=False)
    disease_name = db.Column(db.String(320), nullable=False)
    code = db.Column(db.String(40), nullable=False, default='')

    general_m = db.Column(db.Integer, nullable=False, default=0)
    general_f = db.Column(db.Integer, nullable=False, default=0)
    general_tr = db.Column(db.Integer, nullable=False, default=0)
    general_total = db.Column(db.Integer, nullable=False, default=0)

    emergency_m = db.Column(db.Integer, nullable=False, default=0)
    emergency_f = db.Column(db.Integer, nullable=False, default=0)
    emergency_tr = db.Column(db.Integer, nullable=False, default=0)
    emergency_total = db.Column(db.Integer, nullable=False, default=0)

    ipd_general_m = db.Column(db.Integer, nullable=False, default=0)
    ipd_general_f = db.Column(db.Integer, nullable=False, default=0)
    ipd_general_tr = db.Column(db.Integer, nullable=False, default=0)
    ipd_general_total = db.Column(db.Integer, nullable=False, default=0)

    ipd_emergency_m = db.Column(db.Integer, nullable=False, default=0)
    ipd_emergency_f = db.Column(db.Integer, nullable=False, default=0)
    ipd_emergency_tr = db.Column(db.Integer, nullable=False, default=0)
    ipd_emergency_total = db.Column(db.Integer, nullable=False, default=0)

    overall_m = db.Column(db.Integer, nullable=False, default=0)
    overall_f = db.Column(db.Integer, nullable=False, default=0)
    overall_tr = db.Column(db.Integer, nullable=False, default=0)
    overall_total = db.Column(db.Integer, nullable=False, default=0)

    deaths_m = db.Column(db.Integer, nullable=False, default=0)
    deaths_f = db.Column(db.Integer, nullable=False, default=0)
    deaths_tr = db.Column(db.Integer, nullable=False, default=0)
    deaths_total = db.Column(db.Integer, nullable=False, default=0)

    remarks = db.Column(db.String(220), nullable=False, default='')
    is_custom = db.Column(db.Boolean, nullable=False, default=False)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

    __table_args__ = (
        db.UniqueConstraint('month_year', 'sr_no', name='uq_cbhi_form2_month_sr'),
    )


class CbhiForm2Meta(db.Model):
    __tablename__ = 'cbhi_form2_meta'

    id = db.Column(db.Integer, primary_key=True)
    month_year = db.Column(db.String(20), nullable=False, unique=True, index=True)
    health_establishment = db.Column(db.String(220), nullable=False, default='PRIMARY HEALTH CENTRE, POSSI')
    establishment_phone = db.Column(db.String(60), nullable=False, default='')
    complete_address = db.Column(db.String(320), nullable=False, default='VILL: POSSI, PO: POSSI, DISTRICT-HOSHIARPUR (PUNJAB)')
    district = db.Column(db.String(120), nullable=False, default='HOSHIARPUR (PUNJAB)')
    approving_authority = db.Column(db.String(160), nullable=False, default='')
    authority_designation = db.Column(db.String(160), nullable=False, default='')
    official_email = db.Column(db.String(160), nullable=False, default='')
    official_phone = db.Column(db.String(40), nullable=False, default='')


class UserReportSubmission(db.Model):
    __tablename__ = 'user_report_submissions'

    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=False, index=True)
    report_type = db.Column(db.String(30), nullable=False, index=True)
    month_year = db.Column(db.String(20), nullable=False, index=True)
    total_opd = db.Column(db.Integer, nullable=False, default=0)
    total_ipd = db.Column(db.Integer, nullable=False, default=0)
    total_value = db.Column(db.Integer, nullable=False, default=0)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

    user = db.relationship('User', backref='report_submissions')

    __table_args__ = (
        db.UniqueConstraint('user_id', 'report_type', 'month_year', name='uq_user_report_month_type'),
    )


class UserReportStatus(db.Model):
    __tablename__ = 'user_report_statuses'

    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=False, index=True)
    report_type = db.Column(db.String(30), nullable=False, index=True)
    month_year = db.Column(db.String(20), nullable=False, index=True)
    status = db.Column(db.String(20), nullable=False, default='draft')
    submitted_at = db.Column(db.DateTime, nullable=True)
    reviewed_at = db.Column(db.DateTime, nullable=True)
    reviewed_by = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=True)

    user = db.relationship('User', foreign_keys=[user_id])

    __table_args__ = (
        db.UniqueConstraint('user_id', 'report_type', 'month_year', name='uq_user_report_status_month_type'),
    )


@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


def _normalize_username(value):
    """Strip whitespace and lowercase – every stored username is lowercase."""
    return str(value or '').strip().lower()


def _username_match_query(username):
    """Case-insensitive lookup so legacy uppercase rows are still found."""
    normalized = _normalize_username(username)
    username_expr = db.func.lower(db.func.trim(User.username))
    return User.query.filter(username_expr == normalized)


def _login_candidate_users(username):
    """Return possible login candidates for legacy-renamed usernames.

    Exact lowercase match is preferred. As a compatibility fallback, allow
    usernames that were auto-repaired to "<base>_<id>" format.
    """
    normalized = _normalize_username(username)
    escaped = normalized.replace('\\', '\\\\').replace('%', '\\%').replace('_', '\\_')
    suffix_pattern = f'{escaped}\\_%'
    username_expr = db.func.lower(db.func.trim(User.username))
    return User.query.filter(
        db.or_(
            username_expr == normalized,
            username_expr.like(suffix_pattern, escape='\\'),
        )
    ).order_by(User.created_at.asc()).all()


def _parse_month_year(value):
    try:
        return datetime.strptime((value or '').upper(), '%b-%Y')
    except ValueError:
        return datetime.min


def _normalize_month_year(value, fallback=None):
    default_value = fallback or datetime.utcnow().strftime('%b-%Y')
    raw_value = (value or default_value).strip().upper()
    parsed = _parse_month_year(raw_value)
    if parsed == datetime.min:
        parsed = _parse_month_year(default_value)
    if parsed == datetime.min:
        parsed = datetime.utcnow()
    return parsed.strftime('%b-%Y').upper()


def _user_scoped_month_key(month_year, user_id):
    normalized_month = _normalize_month_year(month_year)
    return f'{normalized_month}__U{user_id}'


# ── Forms ────────────────────────────────────────────────────────────────────

class LoginForm(FlaskForm):
    username = StringField('Username', validators=[DataRequired(), Length(3, 80)])
    password = PasswordField('Password', validators=[DataRequired()])
    submit   = SubmitField('Sign In')


class SelfRegisterForm(FlaskForm):
    username = StringField('Username', validators=[DataRequired(), Length(3, 80)])
    password = PasswordField('Password', validators=[DataRequired(), Length(min=8)])
    confirm  = PasswordField('Confirm Password', validators=[DataRequired(), EqualTo('password', message='Passwords must match.')])
    submit   = SubmitField('Create Account')

    def validate_username(self, field):
        username = _normalize_username(field.data)
        if not username:
            raise ValidationError('Username is required.')
        if any(ch.isspace() for ch in username):
            raise ValidationError('Username cannot contain spaces.')
        if _username_match_query(username).first():
            raise ValidationError('Username already taken.')


class CreateUserForm(FlaskForm):
    username = StringField('Username', validators=[DataRequired(), Length(3, 80)])
    password = PasswordField('Password', validators=[DataRequired(), Length(min=8)])
    confirm  = PasswordField('Confirm Password', validators=[DataRequired(), EqualTo('password', message='Passwords must match.')])
    role     = SelectField('Role', choices=[('sub', 'General User'), ('admin', 'Admin')])
    submit   = SubmitField('Create User')

    def validate_username(self, field):
        username = _normalize_username(field.data)
        if not username:
            raise ValidationError('Username is required.')
        if any(ch.isspace() for ch in username):
            raise ValidationError('Username cannot contain spaces.')
        if _username_match_query(username).first():
            raise ValidationError('Username already taken.')


class ProfileForm(FlaskForm):
    username = StringField('Username', validators=[DataRequired(), Length(3, 80)])
    submit   = SubmitField('Save Changes')

    def validate_username(self, field):
        username = _normalize_username(field.data)
        if not username:
            raise ValidationError('Username is required.')
        if any(ch.isspace() for ch in username):
            raise ValidationError('Username cannot contain spaces.')
        user = _username_match_query(username).first()
        if user and user.id != current_user.id:
            raise ValidationError('Username already taken.')


class ChangePasswordForm(FlaskForm):
    current_password = PasswordField('Current Password', validators=[DataRequired()])
    new_password     = PasswordField('New Password', validators=[DataRequired(), Length(min=8)])
    confirm          = PasswordField('Confirm New Password', validators=[DataRequired(), EqualTo('new_password', message='Passwords must match.')])
    submit           = SubmitField('Update Password')


class EditUserForm(FlaskForm):
    username     = StringField('Username', validators=[DataRequired(), Length(3, 80)])
    role         = SelectField('Role', choices=[('sub', 'General User'), ('admin', 'Admin')])
    new_password = PasswordField('New Password (leave blank to keep current)', validators=[Length(min=0)])
    confirm      = PasswordField('Confirm New Password', validators=[EqualTo('new_password', message='Passwords must match.')])
    submit       = SubmitField('Save Changes')

    def __init__(self, user_id, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._user_id = user_id

    def validate_username(self, field):
        username = _normalize_username(field.data)
        if not username:
            raise ValidationError('Username is required.')
        if any(ch.isspace() for ch in username):
            raise ValidationError('Username cannot contain spaces.')
        user = _username_match_query(username).first()
        if user and user.id != self._user_id:
            raise ValidationError('Username already taken.')

    def validate_new_password(self, field):
        if field.data and len(field.data) < 8:
            raise ValidationError('Password must be at least 8 characters.')


# ── Routes ───────────────────────────────────────────────────────────────────

@app.route('/')
def index():
    return redirect(url_for('dashboard')) if current_user.is_authenticated else render_template('home.html')


@app.route('/about')
def about():
    return render_template('about.html')


@app.route('/help')
def help_page():
    return render_template('help.html')


@app.route('/contact')
def contact():
    return render_template('contact.html')


@app.route('/reports/export/<report_type>/<fmt>')
@login_required
def report_export(report_type, fmt):
    if report_type not in {'hospital_indicator', 'proforma_i', 'proforma_ii', 'cbhi_form1', 'cbhi_form2'}:
        abort(404)
    if fmt not in {'print', 'csv', 'pdf'}:
        abort(404)

    month_year = _normalize_month_year(request.args.get('month_year'))
    payload = _report_export_payload(report_type, month_year, current_user.id)
    if not payload:
        abort(404)

    safe_name = report_type.replace('_', '-')
    if fmt == 'print':
        return render_template(
            'report_export_print.html',
            payload=payload,
            username=current_user.username,
        )

    if fmt == 'csv':
        buffer = StringIO()
        writer = csv.writer(buffer)
        writer.writerow([payload['title']])
        writer.writerow(['Month-Year', payload['month_year']])
        writer.writerow(['User', current_user.username])
        writer.writerow([])
        writer.writerow(payload['headers'])
        writer.writerows(payload['rows'])
        filename = f'{safe_name}_{payload["month_year"]}_{current_user.username}.csv'
        return Response(
            buffer.getvalue(),
            mimetype='text/csv',
            headers={'Content-Disposition': f'attachment; filename={filename}'},
        )

    pdf_bytes = _report_payload_pdf_bytes(payload, current_user.username)
    filename = f'{safe_name}_{payload["month_year"]}_{current_user.username}.pdf'
    return Response(
        pdf_bytes,
        mimetype='application/pdf',
        headers={'Content-Disposition': f'attachment; filename={filename}'},
    )


def _to_non_negative_int(value):
    try:
        return max(0, int(str(value).strip()))
    except (ValueError, TypeError):
        return 0


def _client_ip_address():
    forwarded_for = (request.headers.get('X-Forwarded-For') or '').strip()
    if forwarded_for:
        return forwarded_for.split(',')[0].strip()[:64]
    return (request.remote_addr or '')[:64]


def _record_login_attempt(username, success, reason, user=None):
    try:
        db.session.add(LoginAudit(
            username=(username or '').strip()[:80] or 'unknown',
            user_id=(user.id if user else None),
            success=bool(success),
            reason=(reason or '').strip()[:120],
            ip_address=_client_ip_address(),
            user_agent=(request.user_agent.string if request.user_agent else '')[:260],
        ))
        db.session.commit()
    except Exception:
        db.session.rollback()


def _upsert_report_submission(user_id, month_year, report_type, total_opd=0, total_ipd=0, total_value=0):
    submission = UserReportSubmission.query.filter_by(
        user_id=user_id,
        report_type=report_type,
        month_year=month_year,
    ).first()
    if not submission:
        submission = UserReportSubmission(
            user_id=user_id,
            report_type=report_type,
            month_year=month_year,
        )
        db.session.add(submission)
    submission.total_opd = _to_non_negative_int(total_opd)
    submission.total_ipd = _to_non_negative_int(total_ipd)
    submission.total_value = _to_non_negative_int(total_value)


def _get_or_create_report_status(user_id, month_year, report_type):
    row = UserReportStatus.query.filter_by(
        user_id=user_id,
        month_year=month_year,
        report_type=report_type,
    ).first()
    if not row:
        row = UserReportStatus(
            user_id=user_id,
            month_year=month_year,
            report_type=report_type,
            status='draft',
        )
        db.session.add(row)
    return row


def _set_report_status(user_id, month_year, report_type, status, reviewed_by=None):
    row = _get_or_create_report_status(user_id, month_year, report_type)
    row.status = status
    if status == 'submitted':
        row.submitted_at = datetime.utcnow()
        row.reviewed_at = None
        row.reviewed_by = None
    elif status in ('approved', 'rejected'):
        row.reviewed_at = datetime.utcnow()
        row.reviewed_by = reviewed_by
    elif status == 'draft':
        row.reviewed_at = None
        row.reviewed_by = None
    return row


def _report_export_payload(report_type, month_year, user_id):
    scoped_month_year = _user_scoped_month_key(month_year, user_id)
    report_titles = {
        'hospital_indicator': 'Hospital Indicator Report',
        'proforma_i': 'PROFORMA-I (HPI)',
        'proforma_ii': 'PROFORMA-II (Morbidity)',
        'cbhi_form1': 'CBHI FORM-1',
        'cbhi_form2': 'CBHI FORM-2',
    }

    if report_type == 'hospital_indicator':
        rows = HospitalIndicator.query.filter_by(month_year=scoped_month_year).order_by(HospitalIndicator.indicator_no.asc()).all()
        headers = ['Sr No', 'Indicator', 'OPD', 'IPD']
        data_rows = [[r.indicator_no, r.indicator_name, r.opd_count, r.ipd_count] for r in rows]
    elif report_type == 'proforma_i':
        rows = ProformaHPIRow.query.filter_by(month_year=scoped_month_year).order_by(ProformaHPIRow.id.asc()).all()
        headers = ['Code', 'Indicator', 'Male', 'Female', 'Male Child <14', 'Female Child <14', 'Total', 'Remarks']
        data_rows = [[r.indicator_code, r.indicator_label, r.male, r.female, r.male_child_u14, r.female_child_u14, r.total, r.remarks] for r in rows]
    elif report_type == 'proforma_ii':
        rows = ProformaIIRow.query.filter_by(month_year=scoped_month_year).order_by(ProformaIIRow.sr_no.asc()).all()
        headers = ['Sr No', 'Disease', 'OPD', 'IPD']
        data_rows = [[r.sr_no, r.disease_name, r.opd_count, r.ipd_count] for r in rows]
    elif report_type == 'cbhi_form1':
        rows = CbhiForm1Row.query.filter_by(month_year=scoped_month_year).order_by(CbhiForm1Row.sr_no.asc()).all()
        headers = [
            'Sr No', 'Disease', 'Code',
            'General M', 'General F', 'General TR', 'General Total',
            'Emergency M', 'Emergency F', 'Emergency TR', 'Emergency Total',
            'IPD General M', 'IPD General F', 'IPD General TR', 'IPD General Total',
            'IPD Emergency M', 'IPD Emergency F', 'IPD Emergency TR', 'IPD Emergency Total',
            'Overall M', 'Overall F', 'Overall TR', 'Overall Total',
            'Deaths M', 'Deaths F', 'Deaths TR', 'Deaths Total',
            'Remarks',
        ]
        data_rows = [[
            r.sr_no, r.disease_name, r.code,
            r.general_m, r.general_f, r.general_tr, r.general_total,
            r.emergency_m, r.emergency_f, r.emergency_tr, r.emergency_total,
            r.ipd_general_m, r.ipd_general_f, r.ipd_general_tr, r.ipd_general_total,
            r.ipd_emergency_m, r.ipd_emergency_f, r.ipd_emergency_tr, r.ipd_emergency_total,
            r.overall_m, r.overall_f, r.overall_tr, r.overall_total,
            r.deaths_m, r.deaths_f, r.deaths_tr, r.deaths_total,
            r.remarks,
        ] for r in rows]
    elif report_type == 'cbhi_form2':
        rows = CbhiForm2Row.query.filter_by(month_year=scoped_month_year).order_by(CbhiForm2Row.sr_no.asc()).all()
        headers = [
            'Sr No', 'Disease', 'Code',
            'General M', 'General F', 'General TR', 'General Total',
            'Emergency M', 'Emergency F', 'Emergency TR', 'Emergency Total',
            'IPD General M', 'IPD General F', 'IPD General TR', 'IPD General Total',
            'IPD Emergency M', 'IPD Emergency F', 'IPD Emergency TR', 'IPD Emergency Total',
            'Overall M', 'Overall F', 'Overall TR', 'Overall Total',
            'Deaths M', 'Deaths F', 'Deaths TR', 'Deaths Total',
            'Remarks',
        ]
        data_rows = [[
            r.sr_no, r.disease_name, r.code,
            r.general_m, r.general_f, r.general_tr, r.general_total,
            r.emergency_m, r.emergency_f, r.emergency_tr, r.emergency_total,
            r.ipd_general_m, r.ipd_general_f, r.ipd_general_tr, r.ipd_general_total,
            r.ipd_emergency_m, r.ipd_emergency_f, r.ipd_emergency_tr, r.ipd_emergency_total,
            r.overall_m, r.overall_f, r.overall_tr, r.overall_total,
            r.deaths_m, r.deaths_f, r.deaths_tr, r.deaths_total,
            r.remarks,
        ] for r in rows]
    else:
        return None

    return {
        'title': report_titles[report_type],
        'month_year': month_year,
        'headers': headers,
        'rows': data_rows,
    }


def _report_payload_pdf_bytes(payload, username):
    buffer = BytesIO()
    doc = canvas.Canvas(buffer, pagesize=landscape(A4))
    width, height = landscape(A4)
    left = 20
    top = height - 24
    line_height = 11

    def _new_page():
        doc.showPage()
        doc.setFillColorRGB(0, 0, 0)
        doc.setStrokeColorRGB(0, 0, 0)

    y = top
    doc.setFont('Helvetica-Bold', 12)
    doc.drawString(left, y, f"{payload['title']} - {payload['month_year']}")
    y -= line_height
    doc.setFont('Helvetica', 9)
    doc.drawString(left, y, f'User: {username}')
    y -= (line_height + 2)

    header_line = ' | '.join(str(h) for h in payload['headers'])
    doc.setFont('Helvetica-Bold', 8)
    for part in wrap(header_line, width=175):
        if y < 24:
            _new_page()
            y = top
        doc.drawString(left, y, part)
        y -= line_height

    doc.setFont('Helvetica', 8)
    for row in payload['rows']:
        row_line = ' | '.join(str(v) for v in row)
        wrapped = wrap(row_line, width=175) or ['']
        for part in wrapped:
            if y < 24:
                _new_page()
                y = top
            doc.drawString(left, y, part)
            y -= line_height

    doc.save()
    return buffer.getvalue()


def _report_payload_csv_text(payload, username):
    buffer = StringIO()
    writer = csv.writer(buffer)
    writer.writerow([payload['title']])
    writer.writerow(['Month-Year', payload['month_year']])
    writer.writerow(['User', username])
    writer.writerow([])
    writer.writerow(payload['headers'])
    writer.writerows(payload['rows'])
    return buffer.getvalue()


def _consolidated_report_payload(report_type, month_year):
    """Aggregate row-level data across ALL users for a given month, producing
    a single consolidated payload with summed numeric values."""
    report_titles = {
        'hospital_indicator': 'Hospital Indicator Report (Consolidated)',
        'proforma_i': 'PROFORMA-I HPI (Consolidated)',
        'proforma_ii': 'PROFORMA-II Morbidity (Consolidated)',
        'cbhi_form1': 'CBHI FORM-1 (Consolidated)',
        'cbhi_form2': 'CBHI FORM-2 (Consolidated)',
    }
    if report_type not in report_titles:
        return None

    normalized = _normalize_month_year(month_year)

    # Row data is stored with scoped key like 'MAR-2026__U5'.
    # Use LIKE to match all users' rows for this month, and also match
    # any legacy rows stored with plain month_year.
    scoped_pattern = f'{normalized}__U%'

    # Count how many users submitted this report for this month.
    submissions = UserReportSubmission.query.filter_by(
        report_type=report_type, month_year=normalized
    ).all()
    user_count = len(submissions)

    def _month_filter(model):
        """Return filter matching both scoped (MAR-2026__U5) and plain (MAR-2026) keys."""
        return db.or_(
            model.month_year.like(scoped_pattern),
            model.month_year == normalized,
        )

    if report_type == 'hospital_indicator':
        all_rows = HospitalIndicator.query.filter(_month_filter(HospitalIndicator)).all()
        agg = {}
        for r in all_rows:
            key = r.indicator_no
            if key not in agg:
                agg[key] = {'indicator_no': r.indicator_no, 'indicator_name': r.indicator_name,
                            'opd_count': 0, 'ipd_count': 0}
            agg[key]['opd_count'] += (r.opd_count or 0)
            agg[key]['ipd_count'] += (r.ipd_count or 0)
        headers = ['Sr No', 'Indicator', 'OPD', 'IPD']
        data_rows = [
            [v['indicator_no'], v['indicator_name'], v['opd_count'], v['ipd_count']]
            for v in sorted(agg.values(), key=lambda x: x['indicator_no'])
        ]

    elif report_type == 'proforma_i':
        all_rows = ProformaHPIRow.query.filter(_month_filter(ProformaHPIRow)).all()
        agg = {}
        for r in all_rows:
            key = r.indicator_code
            if key not in agg:
                agg[key] = {'indicator_code': key, 'indicator_label': r.indicator_label,
                            'male': 0, 'female': 0, 'male_child_u14': 0, 'female_child_u14': 0,
                            'total': 0}
            agg[key]['male'] += (r.male or 0)
            agg[key]['female'] += (r.female or 0)
            agg[key]['male_child_u14'] += (r.male_child_u14 or 0)
            agg[key]['female_child_u14'] += (r.female_child_u14 or 0)
            agg[key]['total'] += (r.total or 0)
        headers = ['Code', 'Indicator', 'Male', 'Female', 'Male Child <14', 'Female Child <14', 'Total']
        data_rows = [
            [v['indicator_code'], v['indicator_label'], v['male'], v['female'],
             v['male_child_u14'], v['female_child_u14'], v['total']]
            for v in sorted(agg.values(), key=lambda x: x['indicator_code'])
        ]

    elif report_type == 'proforma_ii':
        all_rows = ProformaIIRow.query.filter(_month_filter(ProformaIIRow)).all()
        agg = {}
        for r in all_rows:
            key = r.sr_no
            if key not in agg:
                agg[key] = {'sr_no': r.sr_no, 'disease_name': r.disease_name,
                            'opd_count': 0, 'ipd_count': 0}
            agg[key]['opd_count'] += (r.opd_count or 0)
            agg[key]['ipd_count'] += (r.ipd_count or 0)
        headers = ['Sr No', 'Disease', 'OPD', 'IPD']
        data_rows = [
            [v['sr_no'], v['disease_name'], v['opd_count'], v['ipd_count']]
            for v in sorted(agg.values(), key=lambda x: x['sr_no'])
        ]

    elif report_type in ('cbhi_form1', 'cbhi_form2'):
        ModelClass = CbhiForm1Row if report_type == 'cbhi_form1' else CbhiForm2Row
        all_rows = ModelClass.query.filter(_month_filter(ModelClass)).all()
        agg = {}
        numeric_fields = [
            'general_m', 'general_f', 'general_tr', 'general_total',
            'emergency_m', 'emergency_f', 'emergency_tr', 'emergency_total',
            'ipd_general_m', 'ipd_general_f', 'ipd_general_tr', 'ipd_general_total',
            'ipd_emergency_m', 'ipd_emergency_f', 'ipd_emergency_tr', 'ipd_emergency_total',
            'overall_m', 'overall_f', 'overall_tr', 'overall_total',
            'deaths_m', 'deaths_f', 'deaths_tr', 'deaths_total',
        ]
        for r in all_rows:
            key = r.sr_no
            if key not in agg:
                agg[key] = {'sr_no': r.sr_no, 'disease_name': r.disease_name,
                            'code': r.code}
                for field in numeric_fields:
                    agg[key][field] = 0
            for field in numeric_fields:
                agg[key][field] += (getattr(r, field, 0) or 0)

        headers = [
            'Sr No', 'Disease', 'Code',
            'General M', 'General F', 'General TR', 'General Total',
            'Emergency M', 'Emergency F', 'Emergency TR', 'Emergency Total',
            'IPD General M', 'IPD General F', 'IPD General TR', 'IPD General Total',
            'IPD Emergency M', 'IPD Emergency F', 'IPD Emergency TR', 'IPD Emergency Total',
            'Overall M', 'Overall F', 'Overall TR', 'Overall Total',
            'Deaths M', 'Deaths F', 'Deaths TR', 'Deaths Total',
        ]
        data_rows = [
            [v['sr_no'], v['disease_name'], v['code']] +
            [v[f] for f in numeric_fields]
            for v in sorted(agg.values(), key=lambda x: x['sr_no'])
        ]
    else:
        return None

    # Log to help diagnose empty results.
    actual_user_count = user_count
    if not actual_user_count and all_rows:
        actual_user_count = len(set(
            r.month_year.split('__U')[-1] for r in all_rows if '__U' in r.month_year
        ))
    app.logger.info(
        '[CONSOLIDATED] report_type=%s month=%s pattern=%s user_count=%d rows_found=%d',
        report_type, normalized, scoped_pattern, actual_user_count, len(data_rows),
    )

    return {
        'title': report_titles[report_type],
        'month_year': normalized,
        'headers': headers,
        'rows': data_rows,
        'user_count': actual_user_count,
    }


def _consolidated_payload_docx_bytes(payload):
    """Generate a Word (.docx) document from a consolidated payload."""
    doc = DocxDocument()
    style = doc.styles['Normal']
    style.font.size = Pt(9)
    style.font.name = 'Arial'

    doc.add_heading(payload['title'], level=1)
    doc.add_paragraph(f"Month-Year: {payload['month_year']}  |  Users Consolidated: {payload.get('user_count', '-')}")

    if not payload['rows']:
        doc.add_paragraph('No data available for selected month.')
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    headers = payload['headers']
    table = doc.add_table(rows=1 + len(payload['rows']), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)

    for row_idx, row_data in enumerate(payload['rows'], start=1):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _admin_consolidated_month_pdf_bytes(month_year):
    report_titles = {
        'proforma_i': 'PROFORMA-I (HPI)',
        'proforma_ii': 'PROFORMA-II (Morbidity)',
        'cbhi_form1': 'CBHI FORM-1',
        'cbhi_form2': 'CBHI FORM-2',
    }
    report_order = ['proforma_i', 'proforma_ii', 'cbhi_form1', 'cbhi_form2']

    submissions = UserReportSubmission.query.filter_by(month_year=month_year).all()
    statuses = UserReportStatus.query.filter_by(month_year=month_year).all()
    status_map = {(s.user_id, s.report_type): s.status for s in statuses}

    grouped = {key: [] for key in report_order}
    for item in submissions:
        if item.report_type in grouped:
            grouped[item.report_type].append(item)

    for key in grouped:
        grouped[key].sort(key=lambda x: (x.user.username.lower() if x.user else '', x.user_id))

    buffer = BytesIO()
    doc = canvas.Canvas(buffer, pagesize=landscape(A4))
    width, height = landscape(A4)
    left = 20
    top = height - 24
    line_height = 11

    def _new_page():
        doc.showPage()
        doc.setFillColorRGB(0, 0, 0)
        doc.setStrokeColorRGB(0, 0, 0)

    y = top
    doc.setFont('Helvetica-Bold', 13)
    doc.drawString(left, y, f'Consolidated All Users Reports - {month_year}')
    y -= line_height
    doc.setFont('Helvetica', 9)
    doc.drawString(left, y, 'Black-and-white export generated from admin consolidated dashboard')
    y -= (line_height + 4)

    for report_type in report_order:
        title = report_titles[report_type]
        rows = grouped[report_type]

        if y < 46:
            _new_page()
            y = top

        doc.setFont('Helvetica-Bold', 11)
        doc.drawString(left, y, title)
        y -= line_height

        doc.setFont('Helvetica-Bold', 8)
        doc.drawString(left, y, 'User | Status | OPD | IPD | Total Value')
        y -= line_height
        doc.setFont('Helvetica', 8)

        if not rows:
            if y < 24:
                _new_page()
                y = top
            doc.drawString(left, y, 'No submissions')
            y -= (line_height + 2)
            continue

        for item in rows:
            if y < 24:
                _new_page()
                y = top
            username = item.user.username if item.user else f'user_{item.user_id}'
            status = status_map.get((item.user_id, report_type), 'draft')
            line = f'{username} | {status} | {int(item.total_opd or 0)} | {int(item.total_ipd or 0)} | {int(item.total_value or 0)}'
            for part in wrap(line, width=175) or ['']:
                if y < 24:
                    _new_page()
                    y = top
                doc.drawString(left, y, part)
                y -= line_height
        y -= 3

    doc.save()
    return buffer.getvalue()


def _manageable_users_query(manager_user):
    if manager_user.is_super_admin:
        return User.query.filter(User.id != manager_user.id)
    if manager_user.role == 'admin':
        return User.query.filter_by(role='sub')
    return User.query.filter(User.id == -1)


def _apply_user_dashboard_filters(query, search='', role_filter='all', status_filter='all'):
    search = (search or '').strip()
    role_filter = (role_filter or 'all').strip()
    status_filter = (status_filter or 'all').strip()

    if search:
        query = query.filter(db.func.lower(User.username).contains(search.lower()))

    if role_filter in ('super_admin', 'admin', 'sub'):
        query = query.filter(User.role == role_filter)

    if status_filter == 'active':
        query = query.filter(db.or_(User.is_active.is_(True), User.is_active.is_(None)))
    elif status_filter == 'inactive':
        query = query.filter(User.is_active.is_(False))

    return query


def _apply_user_dashboard_sort(query, sort_by='created_at', sort_dir='desc'):
    sort_by = (sort_by or 'created_at').strip()
    sort_dir = (sort_dir or 'desc').strip().lower()
    sort_map = {
        'username': db.func.lower(User.username),
        'role': User.role,
        'status': User.is_active,
        'created_at': User.created_at,
    }
    sort_col = sort_map.get(sort_by, User.created_at)
    if sort_dir == 'asc':
        return query.order_by(sort_col.asc())
    return query.order_by(sort_col.desc())


def _load_consolidated_submissions(report_type, month_year, search='', status_filter='all', sort_by='updated_at', sort_dir='desc'):
    if not month_year:
        return []

    submissions = UserReportSubmission.query.filter_by(
        report_type=report_type,
        month_year=month_year,
    ).all()

    status_rows = UserReportStatus.query.filter_by(
        report_type=report_type,
        month_year=month_year,
    ).all()
    status_map = {row.user_id: row.status for row in status_rows}

    for item in submissions:
        item.workflow_status = status_map.get(item.user_id, 'draft')

    search = (search or '').strip().lower()
    if search:
        submissions = [
            item for item in submissions
            if search in ((item.user.username if item.user else 'deleted user').lower())
        ]

    status_filter = (status_filter or 'all').strip().lower()
    if status_filter in ('draft', 'submitted', 'approved', 'rejected'):
        submissions = [item for item in submissions if item.workflow_status == status_filter]

    sort_by = (sort_by or 'updated_at').strip()
    sort_dir = (sort_dir or 'desc').strip().lower()
    reverse = sort_dir != 'asc'

    if sort_by == 'username':
        submissions.sort(key=lambda x: (x.user.username.lower() if x.user else 'zzzz'), reverse=reverse)
    elif sort_by == 'status':
        submissions.sort(key=lambda x: x.workflow_status, reverse=reverse)
    elif sort_by == 'total_opd':
        submissions.sort(key=lambda x: x.total_opd, reverse=reverse)
    elif sort_by == 'total_ipd':
        submissions.sort(key=lambda x: x.total_ipd, reverse=reverse)
    elif sort_by == 'total_value':
        submissions.sort(key=lambda x: x.total_value, reverse=reverse)
    else:
        submissions.sort(key=lambda x: x.updated_at or datetime.min, reverse=reverse)

    return submissions


def _ensure_hospital_indicator_rows(month_year):
    existing_count = HospitalIndicator.query.filter_by(month_year=month_year).count()
    if existing_count == 0:
        for number, name in HOSPITAL_INDICATOR_DEFAULTS:
            db.session.add(HospitalIndicator(
                month_year=month_year,
                indicator_no=number,
                indicator_name=name,
                opd_count=0,
                ipd_count=0,
            ))
        db.session.commit()

    meta = HospitalIndicatorMeta.query.filter_by(month_year=month_year).first()
    if not meta:
        meta = HospitalIndicatorMeta(month_year=month_year, institution_name='PHC POSSI')
        db.session.add(meta)
        db.session.commit()


def _ensure_proforma_hpi_rows(month_year):
    existing_count = ProformaHPIRow.query.filter_by(month_year=month_year).count()
    if existing_count == 0:
        for code, label in PROFORMA_HPI_DEFAULTS:
            db.session.add(ProformaHPIRow(
                month_year=month_year,
                indicator_code=code,
                indicator_label=label,
                male=0,
                female=0,
                male_child_u14=0,
                female_child_u14=0,
                total=0,
                remarks='',
            ))
        db.session.commit()

    meta = ProformaHPIMeta.query.filter_by(month_year=month_year).first()
    if not meta:
        db.session.add(ProformaHPIMeta(month_year=month_year))
        db.session.commit()


def _ensure_proforma_ii_rows(month_year):
    existing_count = ProformaIIRow.query.filter_by(month_year=month_year).count()
    if existing_count == 0:
        for sr_no, disease_name in PROFORMA_II_DEFAULTS:
            db.session.add(ProformaIIRow(
                month_year=month_year,
                sr_no=sr_no,
                disease_name=disease_name,
                opd_count=0,
                ipd_count=0,
            ))
        db.session.commit()

    meta = ProformaIIMeta.query.filter_by(month_year=month_year).first()
    if not meta:
        db.session.add(ProformaIIMeta(month_year=month_year, institution_name='PHC POSSI'))
        db.session.commit()


def _ensure_cbhi_form1_rows(month_year):
    existing_count = CbhiForm1Row.query.filter_by(month_year=month_year).count()
    if existing_count == 0:
        for sr_no, name in CBHI_FORM1_DEFAULTS:
            db.session.add(CbhiForm1Row(
                month_year=month_year,
                sr_no=sr_no,
                disease_name=name,
                code='',
            ))
        db.session.commit()

    meta = CbhiForm1Meta.query.filter_by(month_year=month_year).first()
    if not meta:
        db.session.add(CbhiForm1Meta(month_year=month_year))
        db.session.commit()


def _ensure_cbhi_form2_rows(month_year):
    existing_count = CbhiForm2Row.query.filter_by(month_year=month_year).count()
    if existing_count == 0:
        for sr_no, name in CBHI_FORM2_DEFAULTS:
            db.session.add(CbhiForm2Row(
                month_year=month_year,
                sr_no=sr_no,
                disease_name=name,
                code='',
            ))
        db.session.commit()

    meta = CbhiForm2Meta.query.filter_by(month_year=month_year).first()
    if not meta:
        db.session.add(CbhiForm2Meta(month_year=month_year))
        db.session.commit()


@app.route('/reports/hospital-indicator', methods=['GET', 'POST'])
@login_required
def hospital_indicator_report():
    month_year = _normalize_month_year(request.values.get('month_year'))
    scoped_month_year = _user_scoped_month_key(month_year, current_user.id)
    _ensure_hospital_indicator_rows(scoped_month_year)
    report_type = 'hospital_indicator'

    if request.method == 'POST':
        if not current_user.is_active:
            abort(403)
        action = (request.form.get('action') or '').strip().lower()
        submit_to_admin = request.form.get('submit_to_admin') == '1'

        institution_name = (request.form.get('institution_name') or 'PHC POSSI').strip() or 'PHC POSSI'
        meta = HospitalIndicatorMeta.query.filter_by(month_year=scoped_month_year).first()
        meta.institution_name = institution_name[:150]

        indicators = HospitalIndicator.query.filter_by(month_year=scoped_month_year).order_by(HospitalIndicator.indicator_no.asc()).all()
        if action == 'reset':
            for item in indicators:
                item.opd_count = 0
                item.ipd_count = 0
            _upsert_report_submission(
                user_id=current_user.id,
                month_year=month_year,
                report_type=report_type,
                total_opd=0,
                total_ipd=0,
                total_value=0,
            )
            _set_report_status(current_user.id, month_year, report_type, 'draft')
            db.session.commit()
            flash('Hospital Indicator values reset to zero for this month.', 'success')
            return redirect(url_for('hospital_indicator_report', month_year=month_year))

        for item in indicators:
            item.opd_count = _to_non_negative_int(request.form.get(f'opd_{item.id}', 0))
            item.ipd_count = _to_non_negative_int(request.form.get(f'ipd_{item.id}', 0))

        total_opd = sum(item.opd_count for item in indicators)
        total_ipd = sum(item.ipd_count for item in indicators)
        _upsert_report_submission(
            user_id=current_user.id,
            month_year=month_year,
            report_type=report_type,
            total_opd=total_opd,
            total_ipd=total_ipd,
            total_value=0,
        )
        _get_or_create_report_status(current_user.id, month_year, report_type)
        if submit_to_admin:
            _set_report_status(current_user.id, month_year, report_type, 'submitted')

        db.session.commit()
        flash('Report submitted to admin successfully.' if submit_to_admin else 'Hospital Indicator Report updated successfully.', 'success')
        return redirect(url_for('hospital_indicator_report', month_year=month_year))

    meta = HospitalIndicatorMeta.query.filter_by(month_year=scoped_month_year).first()
    indicators = HospitalIndicator.query.filter_by(month_year=scoped_month_year).order_by(HospitalIndicator.indicator_no.asc()).all()
    total_opd = sum(item.opd_count for item in indicators)
    total_ipd = sum(item.ipd_count for item in indicators)
    status_row = _get_or_create_report_status(current_user.id, month_year, report_type)
    db.session.commit()

    return render_template(
        'hospital_indicator_report.html',
        month_year=month_year,
        institution_name=meta.institution_name,
        indicators=indicators,
        total_opd=total_opd,
        total_ipd=total_ipd,
        submission_status=status_row.status,
    )


@app.route('/reports/proforma-i-hpi', methods=['GET', 'POST'])
@login_required
def proforma_i_hpi_report():
    month_year = _normalize_month_year(request.values.get('month_year'), 'MAR-2026')
    scoped_month_year = _user_scoped_month_key(month_year, current_user.id)
    _ensure_proforma_hpi_rows(scoped_month_year)
    report_type = 'proforma_i'

    if request.method == 'POST':
        if not current_user.is_active:
            abort(403)
        action = (request.form.get('action') or '').strip().lower()
        submit_to_admin = request.form.get('submit_to_admin') == '1'

        meta = ProformaHPIMeta.query.filter_by(month_year=scoped_month_year).first()
        meta.hospital_name = (request.form.get('hospital_name') or '').strip()[:200] or 'COMPILED REPORT OF BLOCK- POSSI'
        meta.district = (request.form.get('district') or '').strip()[:120] or 'HOSHIARPUR'
        meta.sanctioned_beds = (request.form.get('sanctioned_beds') or '').strip()[:60]
        meta.functional_beds = (request.form.get('functional_beds') or '').strip()[:60]
        meta.doctor_incharge = (request.form.get('doctor_incharge') or '').strip()[:120] or 'SMO POSSI'
        meta.note_text = (request.form.get('note_text') or '').strip()[:400] or meta.note_text

        rows = ProformaHPIRow.query.filter_by(month_year=scoped_month_year).all()
        rows.sort(key=lambda r: PROFORMA_HPI_ORDER.get(r.indicator_code, 9999))
        if action == 'reset':
            for row in rows:
                row.male = 0
                row.female = 0
                row.male_child_u14 = 0
                row.female_child_u14 = 0
                row.total = 0
                row.remarks = ''
            _upsert_report_submission(
                user_id=current_user.id,
                month_year=month_year,
                report_type=report_type,
                total_opd=0,
                total_ipd=0,
                total_value=0,
            )
            _set_report_status(current_user.id, month_year, report_type, 'draft')
            db.session.commit()
            flash('PROFORMA-I values reset to zero for this month.', 'success')
            return redirect(url_for('proforma_i_hpi_report', month_year=month_year))

        _single_codes = {'13', '14', '15', '16', '17', '21', '22', '23', '24'}
        for row in rows:
            if row.indicator_code in _single_codes:
                row.male = 0
                row.female = 0
                row.male_child_u14 = 0
                row.female_child_u14 = 0
                row.total = _to_non_negative_int(request.form.get(f'single_{row.id}', 0))
            else:
                row.male = _to_non_negative_int(request.form.get(f'male_{row.id}', 0))
                row.female = _to_non_negative_int(request.form.get(f'female_{row.id}', 0))
                row.male_child_u14 = _to_non_negative_int(request.form.get(f'male_child_{row.id}', 0))
                row.female_child_u14 = _to_non_negative_int(request.form.get(f'female_child_{row.id}', 0))
                row.total = row.male + row.female + row.male_child_u14 + row.female_child_u14
            row.remarks = (request.form.get(f'remarks_{row.id}') or '').strip()[:260]
        
        # Compute auto-sum rows
        _by_code = {r.indicator_code: r for r in rows}

        # Row 1D = Row 1A + Row 1B + Row 1C (column-wise)
        if all(c in _by_code for c in ['1A', '1B', '1C', '1D']):
            _by_code['1D'].male = _by_code['1A'].male + _by_code['1B'].male + _by_code['1C'].male
            _by_code['1D'].female = _by_code['1A'].female + _by_code['1B'].female + _by_code['1C'].female
            _by_code['1D'].male_child_u14 = _by_code['1A'].male_child_u14 + _by_code['1B'].male_child_u14 + _by_code['1C'].male_child_u14
            _by_code['1D'].female_child_u14 = _by_code['1A'].female_child_u14 + _by_code['1B'].female_child_u14 + _by_code['1C'].female_child_u14
            _by_code['1D'].total = _by_code['1D'].male + _by_code['1D'].female + _by_code['1D'].male_child_u14 + _by_code['1D'].female_child_u14
        
        # Row 10 = Row 5 + Row 6 + Row 7 + Row 8 (column-wise)
        if all(c in _by_code for c in ['5', '6', '7', '8', '10']):
            _by_code['10'].male = _by_code['5'].male + _by_code['6'].male + _by_code['7'].male + _by_code['8'].male
            _by_code['10'].female = _by_code['5'].female + _by_code['6'].female + _by_code['7'].female + _by_code['8'].female
            _by_code['10'].male_child_u14 = _by_code['5'].male_child_u14 + _by_code['6'].male_child_u14 + _by_code['7'].male_child_u14 + _by_code['8'].male_child_u14
            _by_code['10'].female_child_u14 = _by_code['5'].female_child_u14 + _by_code['6'].female_child_u14 + _by_code['7'].female_child_u14 + _by_code['8'].female_child_u14
            _by_code['10'].total = _by_code['10'].male + _by_code['10'].female + _by_code['10'].male_child_u14 + _by_code['10'].female_child_u14
        
        # Row 15 = Row 13 + Row 14
        if all(c in _by_code for c in ['13', '14', '15']):
            _by_code['15'].total = _by_code['13'].total + _by_code['14'].total

        _upsert_report_submission(
            user_id=current_user.id,
            month_year=month_year,
            report_type=report_type,
            total_opd=0,
            total_ipd=0,
            total_value=sum(r.total for r in rows),
        )
        _get_or_create_report_status(current_user.id, month_year, report_type)
        if submit_to_admin:
            _set_report_status(current_user.id, month_year, report_type, 'submitted')

        db.session.commit()
        flash('Report submitted to admin successfully.' if submit_to_admin else 'PROFORMA-I HPI report saved successfully.', 'success')
        return redirect(url_for('proforma_i_hpi_report', month_year=month_year))

    meta = ProformaHPIMeta.query.filter_by(month_year=scoped_month_year).first()
    rows = ProformaHPIRow.query.filter_by(month_year=scoped_month_year).all()
    rows.sort(key=lambda r: PROFORMA_HPI_ORDER.get(r.indicator_code, 9999))
    status_row = _get_or_create_report_status(current_user.id, month_year, report_type)
    db.session.commit()

    return render_template(
        'proforma_i_hpi_report.html',
        month_year=month_year,
        meta=meta,
        rows=rows,
        submission_status=status_row.status,
    )


@app.route('/reports/proforma-ii-editable', methods=['GET', 'POST'])
@login_required
def proforma_ii_editable_report():
    month_year = _normalize_month_year(request.values.get('month_year'), 'MAR-2027')
    scoped_month_year = _user_scoped_month_key(month_year, current_user.id)
    _ensure_proforma_ii_rows(scoped_month_year)
    report_type = 'proforma_ii'

    if request.method == 'POST':
        if not current_user.is_active:
            abort(403)
        action = (request.form.get('action') or '').strip().lower()
        submit_to_admin = request.form.get('submit_to_admin') == '1'

        meta = ProformaIIMeta.query.filter_by(month_year=scoped_month_year).first()
        meta.institution_name = (request.form.get('institution_name') or 'PHC POSSI').strip()[:200] or 'PHC POSSI'

        rows = ProformaIIRow.query.filter_by(month_year=scoped_month_year).order_by(ProformaIIRow.sr_no.asc()).all()
        if action == 'reset':
            for row in rows:
                row.opd_count = 0
                row.ipd_count = 0
            _upsert_report_submission(
                user_id=current_user.id,
                month_year=month_year,
                report_type=report_type,
                total_opd=0,
                total_ipd=0,
                total_value=0,
            )
            _set_report_status(current_user.id, month_year, report_type, 'draft')
            db.session.commit()
            flash('PROFORMA-II values reset to zero for this month.', 'success')
            return redirect(url_for('proforma_ii_editable_report', month_year=month_year))

        for row in rows:
            row.opd_count = _to_non_negative_int(request.form.get(f'opd_{row.id}', 0))
            row.ipd_count = _to_non_negative_int(request.form.get(f'ipd_{row.id}', 0))

        total_opd = sum(row.opd_count for row in rows)
        total_ipd = sum(row.ipd_count for row in rows)
        _upsert_report_submission(
            user_id=current_user.id,
            month_year=month_year,
            report_type=report_type,
            total_opd=total_opd,
            total_ipd=total_ipd,
            total_value=0,
        )
        _get_or_create_report_status(current_user.id, month_year, report_type)
        if submit_to_admin:
            _set_report_status(current_user.id, month_year, report_type, 'submitted')

        db.session.commit()
        flash('Report submitted to admin successfully.' if submit_to_admin else 'PROFORMA-II report saved successfully.', 'success')
        return redirect(url_for('proforma_ii_editable_report', month_year=month_year))

    meta = ProformaIIMeta.query.filter_by(month_year=scoped_month_year).first()
    rows = ProformaIIRow.query.filter_by(month_year=scoped_month_year).order_by(ProformaIIRow.sr_no.asc()).all()
    total_opd = sum(row.opd_count for row in rows)
    total_ipd = sum(row.ipd_count for row in rows)
    status_row = _get_or_create_report_status(current_user.id, month_year, report_type)
    db.session.commit()

    return render_template(
        'proforma_ii_editable_report.html',
        month_year=month_year,
        institution_name=meta.institution_name,
        rows=rows,
        total_opd=total_opd,
        total_ipd=total_ipd,
        submission_status=status_row.status,
    )


@app.route('/reports/cbhi-form1', methods=['GET', 'POST'])
@login_required
def cbhi_form1_report():
    month_year = _normalize_month_year(request.values.get('month_year'))
    scoped_month_year = _user_scoped_month_key(month_year, current_user.id)
    _ensure_cbhi_form1_rows(scoped_month_year)
    report_type = 'cbhi_form1'

    if request.method == 'POST':
        if not current_user.is_active:
            abort(403)

        action = (request.form.get('action') or '').strip().lower()
        meta = CbhiForm1Meta.query.filter_by(month_year=scoped_month_year).first()
        meta.health_establishment = (request.form.get('health_establishment') or '').strip()[:220] or meta.health_establishment
        meta.complete_address = (request.form.get('complete_address') or '').strip()[:320] or meta.complete_address
        meta.approving_authority = (request.form.get('approving_authority') or '').strip()[:160]
        meta.authority_designation = (request.form.get('authority_designation') or '').strip()[:160]
        meta.official_email = (request.form.get('official_email') or '').strip()[:160]
        meta.official_phone = (request.form.get('official_phone') or '').strip()[:40]

        if action == 'add_row':
            max_sr = db.session.query(db.func.max(CbhiForm1Row.sr_no)).filter_by(month_year=scoped_month_year).scalar() or 0
            db.session.add(CbhiForm1Row(
                month_year=scoped_month_year,
                sr_no=max_sr + 1,
                disease_name=(request.form.get('new_disease_name') or 'NEW DISEASE').strip()[:320] or 'NEW DISEASE',
                code=(request.form.get('new_disease_code') or '').strip()[:40],
                is_custom=True,
            ))
            db.session.commit()
            flash('New disease row added.', 'success')
            return redirect(url_for('cbhi_form1_report', month_year=month_year))

        if action == 'reset':
            rows = CbhiForm1Row.query.filter_by(month_year=scoped_month_year).order_by(CbhiForm1Row.sr_no.asc()).all()
            for row in rows:
                for prefix in ['general', 'emergency', 'ipd_general', 'ipd_emergency', 'deaths', 'overall']:
                    setattr(row, f'{prefix}_m', 0)
                    setattr(row, f'{prefix}_f', 0)
                    setattr(row, f'{prefix}_tr', 0)
                    setattr(row, f'{prefix}_total', 0)
                row.remarks = ''
            _upsert_report_submission(
                user_id=current_user.id,
                month_year=month_year,
                report_type=report_type,
                total_opd=0,
                total_ipd=0,
                total_value=0,
            )
            _set_report_status(current_user.id, month_year, report_type, 'draft')
            db.session.commit()
            flash('CBHI FORM-1 values reset to zero for this month.', 'success')
            return redirect(url_for('cbhi_form1_report', month_year=month_year))

        rows = CbhiForm1Row.query.filter_by(month_year=scoped_month_year).order_by(CbhiForm1Row.sr_no.asc()).all()
        for row in rows:
            row.disease_name = (request.form.get(f'disease_{row.id}') or row.disease_name).strip()[:320] or row.disease_name
            row.code = (request.form.get(f'code_{row.id}') or '').strip()[:40]

            for prefix in ['general', 'emergency', 'ipd_general', 'ipd_emergency', 'deaths']:
                m = _to_non_negative_int(request.form.get(f'{prefix}_m_{row.id}', 0))
                f = _to_non_negative_int(request.form.get(f'{prefix}_f_{row.id}', 0))
                tr = _to_non_negative_int(request.form.get(f'{prefix}_tr_{row.id}', 0))
                setattr(row, f'{prefix}_m', m)
                setattr(row, f'{prefix}_f', f)
                setattr(row, f'{prefix}_tr', tr)
                setattr(row, f'{prefix}_total', m + f + tr)

            row.overall_m = row.general_m + row.emergency_m
            row.overall_f = row.general_f + row.emergency_f
            row.overall_tr = row.general_tr + row.emergency_tr
            row.overall_total = row.overall_m + row.overall_f + row.overall_tr
            row.remarks = (request.form.get(f'remarks_{row.id}') or '').strip()[:220]

        _upsert_report_submission(
            user_id=current_user.id,
            month_year=month_year,
            report_type=report_type,
            total_opd=sum(r.general_total + r.emergency_total for r in rows),
            total_ipd=sum(r.ipd_general_total + r.ipd_emergency_total for r in rows),
            total_value=sum(r.overall_total for r in rows),
        )
        _get_or_create_report_status(current_user.id, month_year, report_type)

        if action == 'submit':
            _set_report_status(current_user.id, month_year, report_type, 'submitted')
            flash('FORM-1 submitted to admin successfully.', 'success')
        else:
            flash('FORM-1 draft saved successfully.', 'success')

        db.session.commit()
        return redirect(url_for('cbhi_form1_report', month_year=month_year))

    meta = CbhiForm1Meta.query.filter_by(month_year=scoped_month_year).first()
    rows = CbhiForm1Row.query.filter_by(month_year=scoped_month_year).order_by(CbhiForm1Row.sr_no.asc()).all()
    status_row = _get_or_create_report_status(current_user.id, month_year, report_type)
    db.session.commit()

    return render_template(
        'cbhi_form1_report.html',
        month_year=month_year,
        meta=meta,
        rows=rows,
        submission_status=status_row.status,
    )


@app.route('/reports/cbhi-form2', methods=['GET', 'POST'])
@login_required
def cbhi_form2_report():
    month_year = _normalize_month_year(request.values.get('month_year'))
    scoped_month_year = _user_scoped_month_key(month_year, current_user.id)
    _ensure_cbhi_form2_rows(scoped_month_year)
    report_type = 'cbhi_form2'

    if request.method == 'POST':
        if not current_user.is_active:
            abort(403)

        action = (request.form.get('action') or '').strip().lower()
        meta = CbhiForm2Meta.query.filter_by(month_year=scoped_month_year).first()
        meta.health_establishment = (request.form.get('health_establishment') or '').strip()[:220] or meta.health_establishment
        meta.establishment_phone = (request.form.get('establishment_phone') or '').strip()[:60]
        meta.complete_address = (request.form.get('complete_address') or '').strip()[:320] or meta.complete_address
        meta.district = (request.form.get('district') or '').strip()[:120] or meta.district
        meta.approving_authority = (request.form.get('approving_authority') or '').strip()[:160]
        meta.authority_designation = (request.form.get('authority_designation') or '').strip()[:160]
        meta.official_email = (request.form.get('official_email') or '').strip()[:160]
        meta.official_phone = (request.form.get('official_phone') or '').strip()[:40]

        if action == 'add_row':
            max_sr = db.session.query(db.func.max(CbhiForm2Row.sr_no)).filter_by(month_year=scoped_month_year).scalar() or 0
            db.session.add(CbhiForm2Row(
                month_year=scoped_month_year,
                sr_no=max_sr + 1,
                disease_name=(request.form.get('new_disease_name') or 'NEW DISEASE').strip()[:320] or 'NEW DISEASE',
                code=(request.form.get('new_disease_code') or '').strip()[:40],
                is_custom=True,
            ))
            db.session.commit()
            flash('New disease row added.', 'success')
            return redirect(url_for('cbhi_form2_report', month_year=month_year))

        if action == 'reset':
            rows = CbhiForm2Row.query.filter_by(month_year=scoped_month_year).order_by(CbhiForm2Row.sr_no.asc()).all()
            for row in rows:
                for prefix in ['general', 'emergency', 'ipd_general', 'ipd_emergency', 'deaths', 'overall']:
                    setattr(row, f'{prefix}_m', 0)
                    setattr(row, f'{prefix}_f', 0)
                    setattr(row, f'{prefix}_tr', 0)
                    setattr(row, f'{prefix}_total', 0)
                row.remarks = ''
            _upsert_report_submission(
                user_id=current_user.id,
                month_year=month_year,
                report_type=report_type,
                total_opd=0,
                total_ipd=0,
                total_value=0,
            )
            _set_report_status(current_user.id, month_year, report_type, 'draft')
            db.session.commit()
            flash('CBHI FORM-2 values reset to zero for this month.', 'success')
            return redirect(url_for('cbhi_form2_report', month_year=month_year))

        rows = CbhiForm2Row.query.filter_by(month_year=scoped_month_year).order_by(CbhiForm2Row.sr_no.asc()).all()
        for row in rows:
            row.disease_name = (request.form.get(f'disease_{row.id}') or row.disease_name).strip()[:320] or row.disease_name
            row.code = (request.form.get(f'code_{row.id}') or '').strip()[:40]

            for prefix in ['general', 'emergency', 'ipd_general', 'ipd_emergency', 'deaths']:
                m = _to_non_negative_int(request.form.get(f'{prefix}_m_{row.id}', 0))
                f = _to_non_negative_int(request.form.get(f'{prefix}_f_{row.id}', 0))
                tr = _to_non_negative_int(request.form.get(f'{prefix}_tr_{row.id}', 0))
                setattr(row, f'{prefix}_m', m)
                setattr(row, f'{prefix}_f', f)
                setattr(row, f'{prefix}_tr', tr)
                setattr(row, f'{prefix}_total', m + f + tr)

            row.overall_m = row.general_m + row.emergency_m
            row.overall_f = row.general_f + row.emergency_f
            row.overall_tr = row.general_tr + row.emergency_tr
            row.overall_total = row.overall_m + row.overall_f + row.overall_tr
            row.remarks = (request.form.get(f'remarks_{row.id}') or '').strip()[:220]

        _upsert_report_submission(
            user_id=current_user.id,
            month_year=month_year,
            report_type=report_type,
            total_opd=sum(r.general_total + r.emergency_total for r in rows),
            total_ipd=sum(r.ipd_general_total + r.ipd_emergency_total for r in rows),
            total_value=sum(r.overall_total for r in rows),
        )
        _get_or_create_report_status(current_user.id, month_year, report_type)

        if action == 'submit':
            _set_report_status(current_user.id, month_year, report_type, 'submitted')
            flash('FORM-2 submitted to admin successfully.', 'success')
        else:
            flash('FORM-2 draft saved successfully.', 'success')

        db.session.commit()
        return redirect(url_for('cbhi_form2_report', month_year=month_year))

    meta = CbhiForm2Meta.query.filter_by(month_year=scoped_month_year).first()
    rows = CbhiForm2Row.query.filter_by(month_year=scoped_month_year).order_by(CbhiForm2Row.sr_no.asc()).all()
    status_row = _get_or_create_report_status(current_user.id, month_year, report_type)
    db.session.commit()

    return render_template(
        'cbhi_form2_report.html',
        month_year=month_year,
        meta=meta,
        rows=rows,
        submission_status=status_row.status,
    )


@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        if not current_user.is_active:
            logout_user()
            session.clear()
        else:
            return redirect(url_for('dashboard'))
    form = LoginForm()
    if form.validate_on_submit():
        username = _normalize_username(form.username.data)
        users = _login_candidate_users(username)
        if not users:
            _record_login_attempt(username=username, success=False, reason='user_not_found')
            flash('Invalid username or password.', 'danger')
            return render_template('login.html', form=form)

        active_users = [u for u in users if u.is_active is not False]
        inactive_users = [u for u in users if u.is_active is False]

        for user in active_users:
            if user.check_password(form.password.data):
                session.clear()
                login_user(user, remember=True)
                session.permanent = True
                _record_login_attempt(username=user.username, success=True, reason='login_success', user=user)
                next_page = request.args.get('next', '')
                # Guard against open redirect
                if next_page.startswith('/') and not next_page.startswith('//'):
                    return redirect(next_page)
                return redirect(url_for('dashboard'))

        for user in inactive_users:
            if user.check_password(form.password.data):
                _record_login_attempt(username=user.username, success=False, reason='inactive_account', user=user)
                flash('Your account is inactive. Contact admin.', 'danger')
                return render_template('login.html', form=form)

        first_user = users[0]
        _record_login_attempt(username=username, success=False, reason='invalid_password', user=first_user)
        flash('Invalid username or password.', 'danger')
    return render_template('login.html', form=form)


@app.route('/signup', methods=['GET', 'POST'])
def signup():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    form = SelfRegisterForm()
    if form.validate_on_submit():
        new_user = User(
            username=_normalize_username(form.username.data),
            role='sub',
            created_by=None
        )
        new_user.set_password(form.password.data)
        db.session.add(new_user)
        db.session.commit()
        flash('Account created successfully. You can now sign in.', 'success')
        return redirect(url_for('login'))
    return render_template('signup.html', form=form)


@app.route('/logout', methods=['POST'])
@login_required
def logout():
    logout_user()
    session.clear()
    flash('You have been signed out.', 'info')
    response = redirect(url_for('login'))
    # Ensure both Flask session and Flask-Login remember cookies are removed.
    response.delete_cookie(app.config.get('SESSION_COOKIE_NAME', 'session'))
    response.delete_cookie(app.config.get('REMEMBER_COOKIE_NAME', 'remember_token'))
    # Prevent stale authenticated pages from being served from browser cache.
    response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
    response.headers['Pragma'] = 'no-cache'
    response.headers['Expires'] = '0'
    return response


@app.route('/register', methods=['GET', 'POST'])
@login_required
def register():
    if not current_user.is_active or not current_user.is_admin_or_above:
        abort(403)
    form = CreateUserForm()
    # Only super_admin may create other admins
    if not current_user.is_super_admin:
        form.role.choices = [('sub', 'General User')]
    if form.validate_on_submit():
        new_user = User(
            username=_normalize_username(form.username.data),
            role=form.role.data,
            created_by=current_user.id,
            is_active=True,
        )
        new_user.set_password(form.password.data)
        db.session.add(new_user)
        db.session.commit()
        flash(f'User "{new_user.username}" created successfully.', 'success')
        return redirect(url_for('dashboard'))
    return render_template('register.html', form=form)


@app.route('/dashboard')
@login_required
def dashboard():
    scope_query = _manageable_users_query(current_user)

    search = request.args.get('q', '').strip()
    role_filter = request.args.get('role', 'all').strip()
    status_filter = request.args.get('status', 'all').strip()
    sort_by = request.args.get('sort_by', 'created_at').strip()
    sort_dir = request.args.get('sort_dir', 'desc').strip().lower()

    # Normalize filters so stale query params never hide users unexpectedly.
    if current_user.is_super_admin:
        allowed_roles = {'all', 'super_admin', 'admin', 'sub'}
    elif current_user.role == 'admin':
        allowed_roles = {'all', 'sub'}
    else:
        allowed_roles = {'all'}
    if role_filter not in allowed_roles:
        role_filter = 'all'

    if status_filter not in {'all', 'active', 'inactive'}:
        status_filter = 'all'
    if sort_by not in {'created_at', 'username', 'role', 'status'}:
        sort_by = 'created_at'
    if sort_dir not in {'asc', 'desc'}:
        sort_dir = 'desc'

    filtered_query = _apply_user_dashboard_filters(
        scope_query,
        search=search,
        role_filter=role_filter,
        status_filter=status_filter,
    )
    users = _apply_user_dashboard_sort(filtered_query, sort_by=sort_by, sort_dir=sort_dir).all()

    stats = {
        'total': scope_query.count(),
        'active': scope_query.filter(db.or_(User.is_active.is_(True), User.is_active.is_(None))).count(),
        'inactive': scope_query.filter(User.is_active.is_(False)).count(),
        'admins': scope_query.filter(User.role.in_(['admin', 'super_admin'])).count(),
        'general_users': scope_query.filter(User.role == 'sub').count(),
        'filtered': len(users),
    }

    login_audits = []
    if current_user.is_admin_or_above:
        login_audits = LoginAudit.query.order_by(LoginAudit.created_at.desc()).limit(12).all()

    return render_template(
        'dashboard.html',
        users=users,
        search=search,
        role_filter=role_filter,
        status_filter=status_filter,
        sort_by=sort_by,
        sort_dir=sort_dir,
        stats=stats,
        login_audits=login_audits,
    )


@app.route('/admin/db-health')
@login_required
def db_health():
    """Admin diagnostic endpoint — shows database type and user counts."""
    if not current_user.is_admin_or_above:
        abort(403)
    engine = db.engine.url.drivername
    total = db.session.execute(db.text("SELECT COUNT(*) FROM users")).scalar()
    active = db.session.execute(
        db.text("SELECT COUNT(*) FROM users WHERE is_active = TRUE OR is_active IS NULL")
    ).scalar()
    inactive = db.session.execute(
        db.text("SELECT COUNT(*) FROM users WHERE is_active = FALSE")
    ).scalar()
    recent = db.session.execute(
        db.text("SELECT username, created_at, is_active FROM users ORDER BY created_at DESC LIMIT 10")
    ).fetchall()
    info = {
        'engine': engine,
        'is_postgresql': 'postgresql' in engine,
        'total_users': total,
        'active_users': active,
        'inactive_users': inactive,
        'recent_users': [
            {'username': r[0], 'created_at': str(r[1]) if r[1] else '', 'is_active': r[2]}
            for r in recent
        ],
        'report_data_counts': {
            'proforma_hpi_rows': db.session.execute(db.text("SELECT COUNT(*) FROM proforma_hpi_rows")).scalar(),
            'proforma_ii_rows': db.session.execute(db.text("SELECT COUNT(*) FROM proforma_ii_rows")).scalar(),
            'hospital_indicators': db.session.execute(db.text("SELECT COUNT(*) FROM hospital_indicators")).scalar(),
            'user_report_submissions': db.session.execute(db.text("SELECT COUNT(*) FROM user_report_submissions")).scalar(),
        },
        'sample_month_year_keys': {
            'proforma_hpi': [r[0] for r in db.session.execute(db.text("SELECT DISTINCT month_year FROM proforma_hpi_rows LIMIT 10")).fetchall()],
            'proforma_ii': [r[0] for r in db.session.execute(db.text("SELECT DISTINCT month_year FROM proforma_ii_rows LIMIT 10")).fetchall()],
            'submissions': [r[0] for r in db.session.execute(db.text("SELECT DISTINCT month_year FROM user_report_submissions LIMIT 10")).fetchall()],
        },
    }
    return jsonify(info)


@app.route('/admin/login-audit')
@login_required
def login_audit():
    if not current_user.is_admin_or_above:
        abort(403)

    q = request.args.get('q', '').strip().lower()
    result = (request.args.get('result') or 'all').strip().lower()

    query = LoginAudit.query
    if q:
        query = query.filter(db.func.lower(LoginAudit.username).contains(q))
    if result == 'success':
        query = query.filter(LoginAudit.success.is_(True))
    elif result == 'failed':
        query = query.filter(LoginAudit.success.is_(False))

    audits = query.order_by(LoginAudit.created_at.desc()).limit(500).all()
    return render_template('login_audit.html', audits=audits, q=q, result=result)


@app.route('/admin/users/export')
@login_required
def export_users_csv():
    if not current_user.is_admin_or_above:
        abort(403)

    scope_query = _manageable_users_query(current_user)
    search = request.args.get('q', '').strip()
    role_filter = request.args.get('role', 'all').strip()
    status_filter = request.args.get('status', 'all').strip()
    sort_by = request.args.get('sort_by', 'created_at').strip()
    sort_dir = request.args.get('sort_dir', 'desc').strip().lower()

    filtered_query = _apply_user_dashboard_filters(
        scope_query,
        search=search,
        role_filter=role_filter,
        status_filter=status_filter,
    )
    users = _apply_user_dashboard_sort(filtered_query, sort_by=sort_by, sort_dir=sort_dir).all()

    buffer = StringIO()
    writer = csv.writer(buffer)
    writer.writerow(['Username', 'Role', 'Status', 'Created At', 'Created By'])
    for user in users:
        writer.writerow([
            user.username,
            user.role,
            'Active' if user.is_active else 'Inactive',
            user.created_at.strftime('%Y-%m-%d %H:%M:%S') if user.created_at else '',
            user.creator.username if user.creator else '',
        ])

    timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
    filename = f'admin_users_{timestamp}.csv'
    return Response(
        buffer.getvalue(),
        mimetype='text/csv',
        headers={'Content-Disposition': f'attachment; filename={filename}'},
    )


@app.route('/admin/users/import/template')
@login_required
def import_users_template_csv():
    if not current_user.is_admin_or_above:
        abort(403)

    buffer = StringIO()
    writer = csv.writer(buffer)
    writer.writerow(['username', 'password', 'role'])
    writer.writerow(['newuser01', 'StrongPass123', 'sub'])
    writer.writerow(['newuser02', 'StrongPass456', 'sub'])

    timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
    filename = f'user_import_template_{timestamp}.csv'
    return Response(
        buffer.getvalue(),
        mimetype='text/csv',
        headers={'Content-Disposition': f'attachment; filename={filename}'},
    )


@app.route('/admin/users/import', methods=['POST'])
@login_required
def import_users_csv():
    if not current_user.is_admin_or_above:
        abort(403)

    upload = request.files.get('users_csv')
    if not upload or not (upload.filename or '').strip():
        flash('Please choose a CSV file to import users.', 'warning')
        return redirect(url_for('dashboard'))

    if not upload.filename.lower().endswith('.csv'):
        flash('Only CSV files are allowed.', 'danger')
        return redirect(url_for('dashboard'))

    try:
        raw_content = upload.stream.read().decode('utf-8-sig')
    except Exception:
        flash('Unable to read CSV file. Please upload a valid UTF-8 CSV.', 'danger')
        return redirect(url_for('dashboard'))

    reader = csv.DictReader(StringIO(raw_content))
    header_map = {str(h or '').strip().lower() for h in (reader.fieldnames or [])}
    if not {'username', 'password'}.issubset(header_map):
        flash('CSV must include headers: username,password (role is optional).', 'danger')
        return redirect(url_for('dashboard'))

    created_users = []
    seen_usernames = set()
    failures = []

    for line_number, row in enumerate(reader, start=2):
        row_data = {str(k or '').strip().lower(): (v or '').strip() for k, v in row.items()}
        username = _normalize_username(row_data.get('username', ''))
        password = row_data.get('password', '')
        role = (row_data.get('role', 'sub') or 'sub').strip().lower()

        if not username:
            failures.append(f'Line {line_number}: username is required.')
            continue
        if any(ch.isspace() for ch in username):
            failures.append(f'Line {line_number}: username cannot contain spaces.')
            continue
        if len(password) < 8:
            failures.append(f'Line {line_number}: password must be at least 8 characters.')
            continue
        if role not in {'sub', 'admin', 'super_admin'}:
            failures.append(f'Line {line_number}: role must be sub or admin.')
            continue
        if not current_user.is_super_admin and role != 'sub':
            failures.append(f'Line {line_number}: only super admin can create admin users.')
            continue

        if username in seen_usernames:
            failures.append(f'Line {line_number}: duplicate username "{username}" in CSV.')
            continue
        if _username_match_query(username).first():
            failures.append(f'Line {line_number}: username "{username}" already exists.')
            continue

        seen_usernames.add(username)
        user = User(
            username=username,
            role=role,
            created_by=current_user.id,
            is_active=True,
        )
        user.set_password(password)
        created_users.append(user)

    if created_users:
        db.session.add_all(created_users)
        db.session.commit()

        # Post-commit verification: confirm users actually persisted.
        verified = 0
        for u in created_users:
            if User.query.get(u.id) is not None:
                verified += 1
        db_engine = db.engine.url.drivername
        app.logger.info(
            'CSV import: %d created, %d verified in DB (engine=%s)',
            len(created_users), verified, db_engine,
        )
        if verified < len(created_users):
            app.logger.error(
                'CSV import VERIFICATION FAILED: only %d/%d users found after commit (engine=%s)',
                verified, len(created_users), db_engine,
            )

    created_count = len(created_users)
    failed_count = len(failures)
    if created_count:
        flash(f'CSV import complete: {created_count} users created.', 'success')
    if failed_count:
        preview = ' | '.join(failures[:3])
        suffix = ' ...' if failed_count > 3 else ''
        flash(f'CSV import skipped {failed_count} rows. {preview}{suffix}', 'warning')
    if not created_count and not failed_count:
        flash('No user rows found in CSV.', 'info')

    return redirect(url_for('dashboard'))


@app.route('/reports/dashboard')
@login_required
def reports_dashboard():
    if not current_user.is_admin_or_above:
        abort(403)
    submissions = UserReportSubmission.query.all()
    statuses = UserReportStatus.query.all()
    all_months = sorted({row.month_year for row in submissions}, key=_parse_month_year, reverse=True)

    monthly_type_summary = {}
    for row in submissions:
        key = (row.month_year, row.report_type)
        bucket = monthly_type_summary.setdefault(key, {'opd': 0, 'ipd': 0, 'value': 0, 'count': 0})
        bucket['opd'] += int(row.total_opd or 0)
        bucket['ipd'] += int(row.total_ipd or 0)
        bucket['value'] += int(row.total_value or 0)
        bucket['count'] += 1

    report_data = []
    for month in all_months:
        p1_data = monthly_type_summary.get((month, 'proforma_i'), {'value': 0, 'count': 0})
        p2_data = monthly_type_summary.get((month, 'proforma_ii'), {'opd': 0, 'ipd': 0, 'count': 0})
        cbhi1_data = monthly_type_summary.get((month, 'cbhi_form1'), {'value': 0, 'count': 0})
        cbhi2_data = monthly_type_summary.get((month, 'cbhi_form2'), {'value': 0, 'count': 0})

        report_data.append({
            'month_year': month,
            'p1': {'exists': p1_data['count'] > 0, 'total': p1_data['value']},
            'p2': {'exists': p2_data['count'] > 0, 'opd': p2_data['opd'], 'ipd': p2_data['ipd']},
            'cbhi': {'exists': cbhi1_data['count'] > 0, 'total': cbhi1_data['value']},
            'cbhi2': {'exists': cbhi2_data['count'] > 0, 'total': cbhi2_data['value']},
        })

    report_alias = {
        'proforma_i': 'p1',
        'proforma_ii': 'p2',
        'cbhi_form1': 'cbhi',
        'cbhi_form2': 'cbhi2',
    }
    status_map = {
        (row.user_id, row.month_year, row.report_type): row.status
        for row in statuses
    }

    by_user_month = {}
    for row in submissions:
        alias = report_alias.get(row.report_type)
        if not alias:
            continue
        key = (row.user_id, row.month_year)
        bucket = by_user_month.setdefault(key, {
            'month_year': row.month_year,
            'user_id': row.user_id,
            'username': row.user.username if row.user else f'user_{row.user_id}',
            'p1': {'exists': False, 'status': '-', 'total': 0},
            'p2': {'exists': False, 'status': '-', 'opd': 0, 'ipd': 0},
            'cbhi': {'exists': False, 'status': '-', 'total': 0},
            'cbhi2': {'exists': False, 'status': '-', 'total': 0},
        })

        status_value = status_map.get((row.user_id, row.month_year, row.report_type), 'draft')
        if alias == 'p2':
            bucket['p2'] = {
                'exists': True,
                'status': status_value,
                'opd': int(row.total_opd or 0),
                'ipd': int(row.total_ipd or 0),
            }
        else:
            bucket[alias] = {
                'exists': True,
                'status': status_value,
                'total': int(row.total_value or 0),
            }

    user_month_rows = sorted(
        by_user_month.values(),
        key=lambda item: (_parse_month_year(item['month_year']), item['username'].lower()),
        reverse=True,
    )

    return render_template('reports_dashboard.html',
        report_data=report_data,
        user_month_rows=user_month_rows,
        total_months=len(all_months),
        p1_count=len({row.month_year for row in submissions if row.report_type == 'proforma_i'}),
        p2_count=len({row.month_year for row in submissions if row.report_type == 'proforma_ii'}),
        cbhi_count=len({row.month_year for row in submissions if row.report_type == 'cbhi_form1'}),
        cbhi2_count=len({row.month_year for row in submissions if row.report_type == 'cbhi_form2'}),
    )


@app.route('/admin/consolidated-reports')
@login_required
def consolidated_reports():
    if not current_user.is_admin_or_above:
        abort(403)

    report_labels = {
        'proforma_i': 'PROFORMA-I (HPI)',
        'proforma_ii': 'PROFORMA-II (Morbidity)',
        'cbhi_form1': 'CBHI FORM-1',
        'cbhi_form2': 'CBHI FORM-2',
    }
    report_type = request.args.get('report_type', 'proforma_i')
    if report_type not in report_labels:
        report_type = 'proforma_i'

    month_options = [
        m[0] for m in db.session.query(UserReportSubmission.month_year)
        .filter_by(report_type=report_type)
        .distinct()
        .all()
    ]
    month_options.sort(key=_parse_month_year, reverse=True)

    month_year = (request.args.get('month_year') or '').upper().strip()
    if not month_year and month_options:
        month_year = month_options[0]

    search = request.args.get('q', '').strip()
    status_filter = (request.args.get('status') or 'all').strip().lower()
    sort_by = (request.args.get('sort_by') or 'updated_at').strip()
    sort_dir = (request.args.get('sort_dir') or 'desc').strip().lower()

    submissions = _load_consolidated_submissions(
        report_type=report_type,
        month_year=month_year,
        search=search,
        status_filter=status_filter,
        sort_by=sort_by,
        sort_dir=sort_dir,
    )

    consolidated = {
        'submitters': len(submissions),
        'total_opd': sum(item.total_opd for item in submissions),
        'total_ipd': sum(item.total_ipd for item in submissions),
        'total_value': sum(item.total_value for item in submissions),
        'draft_count': sum(1 for item in submissions if item.workflow_status == 'draft'),
        'submitted_count': sum(1 for item in submissions if item.workflow_status == 'submitted'),
        'approved_count': sum(1 for item in submissions if item.workflow_status == 'approved'),
        'rejected_count': sum(1 for item in submissions if item.workflow_status == 'rejected'),
    }

    monthly_summary = db.session.query(
        UserReportSubmission.month_year,
        db.func.count(UserReportSubmission.id),
        db.func.sum(UserReportSubmission.total_opd),
        db.func.sum(UserReportSubmission.total_ipd),
        db.func.sum(UserReportSubmission.total_value),
    ).filter_by(report_type=report_type).group_by(UserReportSubmission.month_year).all()

    monthly_rows = [
        {
            'month_year': row[0],
            'submitters': int(row[1] or 0),
            'total_opd': int(row[2] or 0),
            'total_ipd': int(row[3] or 0),
            'total_value': int(row[4] or 0),
        }
        for row in monthly_summary
    ]
    monthly_rows.sort(key=lambda r: _parse_month_year(r['month_year']), reverse=True)

    user_all_reports_rows = []
    if month_year:
        all_month_submissions = UserReportSubmission.query.filter_by(month_year=month_year).all()
        all_month_statuses = UserReportStatus.query.filter_by(month_year=month_year).all()
        status_map = {
            (row.user_id, row.report_type): row.status
            for row in all_month_statuses
        }
        by_user = {}
        for item in all_month_submissions:
            bucket = by_user.setdefault(item.user_id, {
                'user_id': item.user_id,
                'username': item.user.username if item.user else f'user_{item.user_id}',
                'role': item.user.role if item.user else 'unknown',
                'can_export_package': bool(item.user and current_user.can_manage(item.user)),
                'p1': {'exists': False, 'total': 0, 'status': '-'},
                'p2': {'exists': False, 'opd': 0, 'ipd': 0, 'status': '-'},
                'cbhi1': {'exists': False, 'total': 0, 'status': '-'},
                'cbhi2': {'exists': False, 'total': 0, 'status': '-'},
            })

            if item.report_type == 'proforma_i':
                bucket['p1'] = {
                    'exists': True,
                    'total': int(item.total_value or 0),
                    'status': status_map.get((item.user_id, 'proforma_i'), 'draft'),
                }
            elif item.report_type == 'proforma_ii':
                bucket['p2'] = {
                    'exists': True,
                    'opd': int(item.total_opd or 0),
                    'ipd': int(item.total_ipd or 0),
                    'status': status_map.get((item.user_id, 'proforma_ii'), 'draft'),
                }
            elif item.report_type == 'cbhi_form1':
                bucket['cbhi1'] = {
                    'exists': True,
                    'total': int(item.total_value or 0),
                    'status': status_map.get((item.user_id, 'cbhi_form1'), 'draft'),
                }
            elif item.report_type == 'cbhi_form2':
                bucket['cbhi2'] = {
                    'exists': True,
                    'total': int(item.total_value or 0),
                    'status': status_map.get((item.user_id, 'cbhi_form2'), 'draft'),
                }

        user_all_reports_rows = sorted(by_user.values(), key=lambda r: r['username'].lower())

    return render_template(
        'consolidated_reports.html',
        report_labels=report_labels,
        report_type=report_type,
        month_year=month_year,
        month_options=month_options,
        submissions=submissions,
        consolidated=consolidated,
        monthly_rows=monthly_rows,
        user_all_reports_rows=user_all_reports_rows,
        search=search,
        status_filter=status_filter,
        sort_by=sort_by,
        sort_dir=sort_dir,
    )


@app.route('/admin/consolidated-reports/export')
@login_required
def consolidated_reports_export():
    if not current_user.is_admin_or_above:
        abort(403)

    report_type = request.args.get('report_type', 'proforma_i').strip()
    month_year = (request.args.get('month_year') or '').upper().strip()
    search = request.args.get('q', '').strip()
    status_filter = (request.args.get('status') or 'all').strip().lower()
    sort_by = (request.args.get('sort_by') or 'updated_at').strip()
    sort_dir = (request.args.get('sort_dir') or 'desc').strip().lower()

    if report_type not in ('proforma_i', 'proforma_ii', 'cbhi_form1', 'cbhi_form2'):
        report_type = 'proforma_i'

    submissions = _load_consolidated_submissions(
        report_type=report_type,
        month_year=month_year,
        search=search,
        status_filter=status_filter,
        sort_by=sort_by,
        sort_dir=sort_dir,
    )

    buffer = StringIO()
    writer = csv.writer(buffer)
    writer.writerow(['User', 'Role', 'Report Type', 'Month', 'OPD', 'IPD', 'Total Value', 'Status', 'Updated At'])
    for item in submissions:
        writer.writerow([
            item.user.username if item.user else 'Deleted User',
            item.user.role if item.user else 'unknown',
            report_type,
            item.month_year,
            item.total_opd,
            item.total_ipd,
            item.total_value,
            item.workflow_status,
            item.updated_at.strftime('%Y-%m-%d %H:%M:%S') if item.updated_at else '',
        ])

    timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
    filename = f'consolidated_{report_type}_{month_year or "all"}_{timestamp}.csv'
    return Response(
        buffer.getvalue(),
        mimetype='text/csv',
        headers={'Content-Disposition': f'attachment; filename={filename}'},
    )


@app.route('/admin/consolidated-reports/export-all-pdf')
@login_required
def consolidated_reports_export_all_pdf():
    if not current_user.is_admin_or_above:
        abort(403)

    month_year = (request.args.get('month_year') or '').upper().strip()
    if not month_year:
        row = db.session.query(UserReportSubmission.month_year).distinct().all()
        months = sorted([m[0] for m in row], key=_parse_month_year, reverse=True)
        month_year = months[0] if months else ''

    if not month_year:
        flash('No consolidated month found for PDF export.', 'warning')
        return redirect(url_for('consolidated_reports'))

    pdf_bytes = _admin_consolidated_month_pdf_bytes(month_year)
    filename = f'consolidated_all_users_{month_year}.pdf'
    return Response(
        pdf_bytes,
        mimetype='application/pdf',
        headers={'Content-Disposition': f'attachment; filename={filename}'},
    )


@app.route('/admin/consolidated-reports/export-user-package')
@login_required
def consolidated_reports_export_user_package():
    if not current_user.is_admin_or_above:
        abort(403)

    user_id = _to_non_negative_int(request.args.get('user_id'))
    month_year = _normalize_month_year(request.args.get('month_year'))
    target_user = User.query.get_or_404(user_id)
    if not current_user.can_manage(target_user):
        abort(403)

    report_types = ['hospital_indicator', 'proforma_i', 'proforma_ii', 'cbhi_form1', 'cbhi_form2']
    zip_buffer = BytesIO()

    with zipfile.ZipFile(zip_buffer, mode='w', compression=zipfile.ZIP_DEFLATED) as zf:
        manifest = [
            f'User: {target_user.username}',
            f'Month-Year: {month_year}',
            f'Generated At (UTC): {datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")}',
            '',
            'Included files per report: B/W PDF + Excel-compatible CSV',
        ]
        zf.writestr('README.txt', '\n'.join(manifest))

        for report_type in report_types:
            payload = _report_export_payload(report_type, month_year, target_user.id)
            if not payload:
                continue
            base_name = f"{report_type.replace('_', '-')}_{month_year}_{target_user.username}"
            zf.writestr(f'{base_name}.csv', _report_payload_csv_text(payload, target_user.username))
            zf.writestr(f'{base_name}.pdf', _report_payload_pdf_bytes(payload, target_user.username))

    filename = f'user_reports_package_{target_user.username}_{month_year}.zip'
    return Response(
        zip_buffer.getvalue(),
        mimetype='application/zip',
        headers={'Content-Disposition': f'attachment; filename={filename}'},
    )


@app.route('/admin/consolidated-proforma')
@login_required
def consolidated_proforma_view():
    """Read-only consolidated proforma showing aggregated row-level data across
    ALL users for a given month. Includes print/PDF/Excel/Word export buttons."""
    if not current_user.is_admin_or_above:
        abort(403)

    report_type = (request.args.get('report_type') or 'proforma_i').strip()
    if report_type not in ('hospital_indicator', 'proforma_i', 'proforma_ii', 'cbhi_form1', 'cbhi_form2'):
        report_type = 'proforma_i'
    month_year = (request.args.get('month_year') or '').upper().strip()
    if not month_year:
        flash('No month specified.', 'warning')
        return redirect(url_for('consolidated_reports'))

    payload = _consolidated_report_payload(report_type, month_year)
    if not payload:
        flash('Invalid report type.', 'danger')
        return redirect(url_for('consolidated_reports'))

    report_labels = {
        'hospital_indicator': 'Hospital Indicator Report',
        'proforma_i': 'PROFORMA-I (HPI)',
        'proforma_ii': 'PROFORMA-II (Morbidity)',
        'cbhi_form1': 'CBHI FORM-1',
        'cbhi_form2': 'CBHI FORM-2',
    }

    return render_template(
        'consolidated_proforma_view.html',
        payload=payload,
        report_type=report_type,
        month_year=month_year,
        report_labels=report_labels,
    )


@app.route('/admin/consolidated-proforma/export/<fmt>')
@login_required
def consolidated_proforma_export(fmt):
    """Export consolidated proforma as print / pdf / csv (Excel) / docx (Word)."""
    if not current_user.is_admin_or_above:
        abort(403)
    if fmt not in ('print', 'pdf', 'csv', 'docx'):
        abort(404)

    report_type = (request.args.get('report_type') or 'proforma_i').strip()
    if report_type not in ('hospital_indicator', 'proforma_i', 'proforma_ii', 'cbhi_form1', 'cbhi_form2'):
        abort(404)
    month_year = (request.args.get('month_year') or '').upper().strip()
    if not month_year:
        abort(404)

    payload = _consolidated_report_payload(report_type, month_year)
    if not payload:
        abort(404)

    safe_name = f"consolidated_{report_type.replace('_', '-')}_{month_year}"

    if fmt == 'print':
        return render_template(
            'report_export_print.html',
            payload=payload,
            username=f'Consolidated ({payload.get("user_count", 0)} users)',
        )

    if fmt == 'csv':
        buffer = StringIO()
        writer = csv.writer(buffer)
        writer.writerow([payload['title']])
        writer.writerow(['Month-Year', payload['month_year']])
        writer.writerow(['Scope', f'Consolidated ({payload.get("user_count", 0)} users)'])
        writer.writerow([])
        writer.writerow(payload['headers'])
        writer.writerows(payload['rows'])
        return Response(
            buffer.getvalue(),
            mimetype='text/csv',
            headers={'Content-Disposition': f'attachment; filename={safe_name}.csv'},
        )

    if fmt == 'pdf':
        pdf_bytes = _report_payload_pdf_bytes(
            payload,
            f'Consolidated ({payload.get("user_count", 0)} users)',
        )
        return Response(
            pdf_bytes,
            mimetype='application/pdf',
            headers={'Content-Disposition': f'attachment; filename={safe_name}.pdf'},
        )

    # fmt == 'docx'
    docx_bytes = _consolidated_payload_docx_bytes(payload)
    return Response(
        docx_bytes,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename={safe_name}.docx'},
    )


@app.route('/admin/consolidated-reports/status', methods=['POST'])
@login_required
def consolidated_reports_status_update():
    if not current_user.is_admin_or_above:
        abort(403)

    user_id = _to_non_negative_int(request.form.get('user_id'))
    report_type = (request.form.get('report_type') or '').strip()
    month_year = (request.form.get('month_year') or '').strip().upper()
    action = (request.form.get('action') or '').strip().lower()
    next_report_type = (request.form.get('next_report_type') or report_type).strip()
    next_month_year = (request.form.get('next_month_year') or month_year).strip().upper()
    next_search = (request.form.get('next_q') or '').strip()
    next_status = (request.form.get('next_status') or 'all').strip().lower()
    next_sort_by = (request.form.get('next_sort_by') or 'updated_at').strip()
    next_sort_dir = (request.form.get('next_sort_dir') or 'desc').strip().lower()

    if report_type not in ('hospital_indicator', 'proforma_i', 'proforma_ii', 'cbhi_form1', 'cbhi_form2') or action not in ('approve', 'reject', 'reset', 'submit'):
        flash('Invalid status update request.', 'danger')
        return redirect(url_for(
            'consolidated_reports',
            report_type=next_report_type,
            month_year=next_month_year,
            q=next_search,
            status=next_status,
            sort_by=next_sort_by,
            sort_dir=next_sort_dir,
        ))

    if action == 'approve':
        _set_report_status(user_id, month_year, report_type, 'approved', reviewed_by=current_user.id)
    elif action == 'reject':
        _set_report_status(user_id, month_year, report_type, 'rejected', reviewed_by=current_user.id)
    elif action == 'submit':
        _set_report_status(user_id, month_year, report_type, 'submitted')
    else:
        _set_report_status(user_id, month_year, report_type, 'draft')

    db.session.commit()
    flash('Submission status updated.', 'success')
    return redirect(url_for(
        'consolidated_reports',
        report_type=next_report_type,
        month_year=next_month_year,
        q=next_search,
        status=next_status,
        sort_by=next_sort_by,
        sort_dir=next_sort_dir,
    ))


@app.route('/admin/consolidated-reports/bulk-status', methods=['POST'])
@login_required
def consolidated_reports_bulk_status_update():
    if not current_user.is_admin_or_above:
        abort(403)

    report_type = (request.form.get('report_type') or '').strip()
    month_year = (request.form.get('month_year') or '').strip().upper()
    action = (request.form.get('action') or '').strip().lower()
    next_search = (request.form.get('next_q') or '').strip()
    next_status = (request.form.get('next_status') or 'all').strip().lower()
    next_sort_by = (request.form.get('next_sort_by') or 'updated_at').strip()
    next_sort_dir = (request.form.get('next_sort_dir') or 'desc').strip().lower()

    user_ids = []
    for value in request.form.getlist('user_ids'):
        parsed = _to_non_negative_int(value)
        if parsed > 0:
            user_ids.append(parsed)
    user_ids = sorted(set(user_ids))

    if report_type not in ('proforma_i', 'proforma_ii', 'cbhi_form1', 'cbhi_form2') or not month_year:
        flash('Invalid bulk update request.', 'danger')
        return redirect(url_for(
            'consolidated_reports',
            report_type=report_type,
            month_year=month_year,
            q=next_search,
            status=next_status,
            sort_by=next_sort_by,
            sort_dir=next_sort_dir,
        ))

    if action not in ('approve', 'reject', 'reset'):
        flash('Choose a valid bulk action.', 'warning')
        return redirect(url_for(
            'consolidated_reports',
            report_type=report_type,
            month_year=month_year,
            q=next_search,
            status=next_status,
            sort_by=next_sort_by,
            sort_dir=next_sort_dir,
        ))

    if not user_ids:
        flash('Select at least one submission for bulk action.', 'warning')
        return redirect(url_for(
            'consolidated_reports',
            report_type=report_type,
            month_year=month_year,
            q=next_search,
            status=next_status,
            sort_by=next_sort_by,
            sort_dir=next_sort_dir,
        ))

    submissions = UserReportSubmission.query.filter(
        UserReportSubmission.report_type == report_type,
        UserReportSubmission.month_year == month_year,
        UserReportSubmission.user_id.in_(user_ids),
    ).all()

    if not submissions:
        flash('No matching submissions were found for selected users.', 'warning')
        return redirect(url_for(
            'consolidated_reports',
            report_type=report_type,
            month_year=month_year,
            q=next_search,
            status=next_status,
            sort_by=next_sort_by,
            sort_dir=next_sort_dir,
        ))

    updated = 0
    for submission in submissions:
        if action == 'approve':
            _set_report_status(submission.user_id, month_year, report_type, 'approved', reviewed_by=current_user.id)
        elif action == 'reject':
            _set_report_status(submission.user_id, month_year, report_type, 'rejected', reviewed_by=current_user.id)
        else:
            _set_report_status(submission.user_id, month_year, report_type, 'draft')
        updated += 1

    db.session.commit()
    flash(f'Bulk status update applied to {updated} submission(s).', 'success')
    return redirect(url_for(
        'consolidated_reports',
        report_type=report_type,
        month_year=month_year,
        q=next_search,
        status=next_status,
        sort_by=next_sort_by,
        sort_dir=next_sort_dir,
    ))


@app.route('/admin/consolidated-reports/delete', methods=['POST'])
@login_required
def consolidated_reports_delete():
    if not current_user.is_admin_or_above:
        abort(403)

    user_id = _to_non_negative_int(request.form.get('user_id'))
    report_type = (request.form.get('report_type') or '').strip()
    month_year = (request.form.get('month_year') or '').strip().upper()
    next_report_type = (request.form.get('next_report_type') or report_type).strip()
    next_month_year = (request.form.get('next_month_year') or month_year).strip().upper()
    next_search = (request.form.get('next_q') or '').strip()
    next_status = (request.form.get('next_status') or 'all').strip().lower()
    next_sort_by = (request.form.get('next_sort_by') or 'updated_at').strip()
    next_sort_dir = (request.form.get('next_sort_dir') or 'desc').strip().lower()

    allowed_report_types = ('hospital_indicator', 'proforma_i', 'proforma_ii', 'cbhi_form1', 'cbhi_form2')
    if report_type not in allowed_report_types or not user_id or not month_year:
        flash('Invalid delete request.', 'danger')
        return redirect(url_for(
            'consolidated_reports',
            report_type=next_report_type,
            month_year=next_month_year,
            q=next_search,
            status=next_status,
            sort_by=next_sort_by,
            sort_dir=next_sort_dir,
        ))

    submission = UserReportSubmission.query.filter_by(
        user_id=user_id,
        report_type=report_type,
        month_year=month_year,
    ).first()
    status_row = UserReportStatus.query.filter_by(
        user_id=user_id,
        report_type=report_type,
        month_year=month_year,
    ).first()

    if not submission and not status_row:
        flash('No matching submission found to delete.', 'warning')
        return redirect(url_for(
            'consolidated_reports',
            report_type=next_report_type,
            month_year=next_month_year,
            q=next_search,
            status=next_status,
            sort_by=next_sort_by,
            sort_dir=next_sort_dir,
        ))

    if submission:
        db.session.delete(submission)
    if status_row:
        db.session.delete(status_row)

    db.session.commit()
    flash('Report submission deleted successfully.', 'success')
    return redirect(url_for(
        'consolidated_reports',
        report_type=next_report_type,
        month_year=next_month_year,
        q=next_search,
        status=next_status,
        sort_by=next_sort_by,
        sort_dir=next_sort_dir,
    ))


@app.route('/profile', methods=['GET', 'POST'])
@login_required
def profile():
    form = ProfileForm(obj=current_user)
    if form.validate_on_submit():
        current_user.username = _normalize_username(form.username.data)
        db.session.commit()
        flash('Profile updated.', 'success')
        return redirect(url_for('profile'))
    return render_template('profile.html', form=form)


@app.route('/change-password', methods=['GET', 'POST'])
@login_required
def change_password():
    form = ChangePasswordForm()
    if form.validate_on_submit():
        if not current_user.check_password(form.current_password.data):
            flash('Current password is incorrect.', 'danger')
        else:
            current_user.set_password(form.new_password.data)
            db.session.commit()
            flash('Password updated successfully.', 'success')
            return redirect(url_for('profile'))
    return render_template('change_password.html', form=form)


@app.route('/user/<int:user_id>/toggle', methods=['POST'])
@login_required
def toggle_user(user_id):
    if not current_user.is_admin_or_above:
        abort(403)
    user = User.query.get_or_404(user_id)
    if not current_user.can_manage(user):
        abort(403)
    user.is_active = not user.is_active
    db.session.commit()
    status = 'activated' if user.is_active else 'deactivated'
    flash(f'User "{user.username}" {status}.', 'success')
    return redirect(url_for('dashboard'))


@app.route('/user/<int:user_id>/delete', methods=['POST'])
@login_required
def delete_user(user_id):
    if not current_user.is_admin_or_above:
        abort(403)
    user = User.query.get_or_404(user_id)
    if not current_user.can_manage(user):
        abort(403)
    username = user.username
    db.session.delete(user)
    db.session.commit()
    flash(f'User "{username}" deleted.', 'success')
    return redirect(url_for('dashboard'))


@app.route('/user/<int:user_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_user(user_id):
    if not current_user.is_admin_or_above:
        abort(403)
    user = User.query.get_or_404(user_id)
    if not current_user.can_manage(user):
        abort(403)
    form = EditUserForm(user_id=user.id, obj=user)
    if not current_user.is_super_admin:
        form.role.choices = [('sub', 'General User')]
    if form.validate_on_submit():
        user.username = _normalize_username(form.username.data)
        user.role = form.role.data
        if form.new_password.data:
            user.set_password(form.new_password.data)
        db.session.commit()
        flash(f'User "{user.username}" updated successfully.', 'success')
        return redirect(url_for('dashboard'))
    return render_template('edit_user.html', form=form, user=user)


@app.route('/forgot-password')
def forgot_password():
    return render_template('forgot_password.html')


# ── Error Handlers ────────────────────────────────────────────────────────────

@app.errorhandler(403)
def forbidden(e):
    return render_template('404.html', code=403, title='Forbidden',
                           message='You do not have permission to access this page.'), 403


@app.errorhandler(404)
def not_found(e):
    return render_template('404.html', code=404, title='Page Not Found',
                           message='The page you are looking for does not exist.'), 404


@app.errorhandler(500)
def server_error(e):
    db.session.rollback()
    return render_template('500.html'), 500


# ── Startup ───────────────────────────────────────────────────────────────────

def seed_admin():
    db.create_all()

    # ── Startup diagnostic log ───────────────────────────────────────────
    db_engine = db.engine.url.drivername
    total_users = db.session.execute(db.text("SELECT COUNT(*) FROM users")).scalar()
    print(f'[STARTUP] Database engine: {db_engine}')
    print(f'[STARTUP] Total users in database: {total_users}')
    if 'sqlite' in db_engine:
        print('[STARTUP] WARNING: Running on SQLite – data will be lost on restart!')

    # Startup migrations:
    # 1. Fix NULL is_active rows → active.
    # 2. Ensure all admin-created accounts remain active.
    # 3. Enforce case-insensitive unique index when safe.
    #
    # IMPORTANT: Do not auto-rename usernames at startup. That can change
    # issued credentials and cause repeated "user id failed" complaints.
    try:
        db.session.execute(db.text("UPDATE users SET is_active = TRUE WHERE is_active IS NULL"))
        db.session.execute(
            db.text("UPDATE users SET is_active = TRUE WHERE created_by IS NOT NULL")
        )

        # Create index only when there are no case-insensitive duplicates.
        duplicates = db.session.execute(db.text(
            """
            SELECT LOWER(TRIM(username)) AS uname, COUNT(*) AS cnt
            FROM users
            GROUP BY LOWER(TRIM(username))
            HAVING COUNT(*) > 1
            """
        )).fetchall()
        if not duplicates:
            db.session.execute(
                db.text("CREATE UNIQUE INDEX IF NOT EXISTS ux_users_username_lower ON users (LOWER(TRIM(username)))")
            )
        db.session.commit()
    except Exception:
        db.session.rollback()
    if not _username_match_query('admin').first():
        admin = User(username='admin', role='super_admin')
        admin.set_password('admin123')
        db.session.add(admin)
        db.session.commit()
        print('Default admin created  →  admin / admin123')


with app.app_context():
    seed_admin()

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT', 5000)), debug=False)
